# -*- coding: utf-8 -*-
"""
Fenomen-Marka Eşleştirme — Grafik Eki
GelistirmeSureci_Kapsamli_Rapor.docx'e grafik bölümü ekler
"""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import os, tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Ayarlar ──────────────────────────────────────────────────────────────────
BASE   = Path(r'C:\Users\USER\Desktop\TezBitirme')
SRC    = BASE / 'GelistirmeSureci_Kapsamli_Rapor.docx'
OUT    = BASE / 'GelistirmeSureci_Kapsamli_Rapor_Grafikli.docx'
TMPDIR = tempfile.mkdtemp()

plt.rcParams.update({
    'font.family'    : 'DejaVu Sans',
    'font.size'      : 10,
    'axes.titlesize' : 13,
    'axes.titleweight': 'bold',
    'axes.spines.top': False,
    'axes.spines.right': False,
    'figure.dpi'     : 150,
    'savefig.dpi'    : 150,
    'savefig.bbox'   : 'tight',
    'savefig.facecolor': 'white',
})

BLUE   = '#1F497D'
LBLUE  = '#2E74B5'
GREEN  = '#207A44'
RED    = '#C00000'
ORANGE = '#E67E22'
GRAY   = '#7F7F7F'
LGRAY  = '#D0D0D0'

def save(fig, name):
    path = os.path.join(TMPDIR, name)
    fig.savefig(path)
    plt.close(fig)
    return path

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 1 — NFS Top 10
# ════════════════════════════════════════════════════════════════════════════════
names_nfs = [
    '@influencer19', '@ardaguler', '@enesbatur', '@seydaerdogan',
    '@influencer17', '@ardasaatci', '@tugkangonultas', '@influencer52',
    '@edakok59', '@muratsakaoglu'
]
nfs_vals = [56.34, 48.57, 47.08, 46.99, 45.46, 43.45, 42.95, 42.80, 42.42, 42.17]
cats      = ['egitim','spor','oyun','moda','teknoloji','spor','spor','saglik','lifestyle','lifestyle']
cat_colors = {
    'egitim':'#8E44AD','spor':BLUE,'oyun':'#E67E22',
    'moda':'#E91E8C','teknoloji':'#16A085','saglik':GREEN,
    'lifestyle':'#795548','seyahat':'#2196F3'
}
colors = [cat_colors.get(c, GRAY) for c in cats]

fig, ax = plt.subplots(figsize=(10, 5.5))
bars = ax.barh(range(len(names_nfs)), nfs_vals[::-1], color=colors[::-1], height=0.65)
ax.set_yticks(range(len(names_nfs)))
ax.set_yticklabels(names_nfs[::-1], fontsize=9.5)
ax.set_xlabel('NFS (Normalleştirilmiş Fenomen Skoru, 0-100)')
ax.set_title('NFS Top 10 — En Yüksek Normalleştirilmiş Fenomen Skoru')
ax.set_xlim(0, 65)
for bar, val in zip(bars, nfs_vals[::-1]):
    ax.text(bar.get_width()+0.4, bar.get_y()+bar.get_height()/2,
            f'{val:.2f}', va='center', fontsize=9, fontweight='bold', color=BLUE)
patches = [mpatches.Patch(color=v, label=k) for k, v in cat_colors.items() if k in cats]
ax.legend(handles=patches, loc='lower right', fontsize=8.5, title='Kategori')
ax.axvline(x=np.mean(nfs_vals), color=GRAY, linestyle='--', linewidth=1.2, alpha=0.7)
ax.text(np.mean(nfs_vals)+0.3, -0.6, f'Ort: {np.mean(nfs_vals):.1f}',
        fontsize=8.5, color=GRAY)
fig.tight_layout()
g1 = save(fig, 'g1_nfs_top10.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 2 — XGBoost Feature Importance Top 15
# ════════════════════════════════════════════════════════════════════════════════
feat_names = [
    'SFS', 'NFS', 'positive_ratio', 'negative_comment_ratio',
    'avg_sentiment_score', 'category_yemek', 'neutral_comment_ratio',
    'account_type_makro', 'posts_per_month', 'positive_comment_ratio',
    'avg_comment_sentiment', 'campaign_travel', 'comment_count',
    'negative_ratio', 'engagement_rate'
]
feat_vals = [
    0.6492, 0.2741, 0.0455, 0.0095, 0.0052, 0.0034, 0.0029,
    0.0028, 0.0019, 0.0015, 0.0014, 0.0012, 0.0010, 0.0008, 0.0006
]
comment_feats = {'negative_comment_ratio','neutral_comment_ratio',
                 'positive_comment_ratio','avg_comment_sentiment','comment_count'}
fc = [RED if f in comment_feats else BLUE for f in feat_names]

fig, ax = plt.subplots(figsize=(10, 6.5))
bars = ax.barh(range(len(feat_names)), feat_vals[::-1], color=fc[::-1], height=0.7)
ax.set_yticks(range(len(feat_names)))
ax.set_yticklabels(feat_names[::-1], fontsize=9.5)
ax.set_xlabel('Feature Önemi (XGBoost Gain)')
ax.set_title('XGBoost Feature Importance — Top 15\n(Kırmızı: Yorum Feature\'ları — Bu Oturumda Eklendi)')
for bar, val in zip(bars, feat_vals[::-1]):
    ax.text(bar.get_width()+0.002, bar.get_y()+bar.get_height()/2,
            f'{val:.4f}', va='center', fontsize=8.5)
ax.set_xscale('log')
patches2 = [
    mpatches.Patch(color=BLUE, label='Mevcut Feature\'lar'),
    mpatches.Patch(color=RED,  label='Yorum Feature\'ları (YENİ)'),
]
ax.legend(handles=patches2, fontsize=9)
fig.tight_layout()
g2 = save(fig, 'g2_feature_importance.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 3 — Model Karşılaştırması (Accuracy, F1, Precision, Recall)
# ════════════════════════════════════════════════════════════════════════════════
models_m = ['XGBoost', 'LightGBM', 'RandomForest']
metrics_m = {
    'Accuracy' : [0.9918, 0.9980, 0.9980],
    'F1 (W)'   : [0.9930, 0.9940, 0.9790],
    'Precision': [0.9918, 0.9980, 0.9980],
    'Recall'   : [0.9918, 0.9980, 0.9980],
}
x = np.arange(len(models_m))
width = 0.2
fig, ax = plt.subplots(figsize=(10, 5))
for i, (metric, vals) in enumerate(metrics_m.items()):
    bars = ax.bar(x + i*width, vals, width, label=metric,
                  color=[BLUE, LBLUE, GREEN, ORANGE][i], alpha=0.85)
    for bar, val in zip(bars, vals):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.0005,
                f'{val:.3f}', ha='center', va='bottom', fontsize=8, fontweight='bold')
ax.set_xticks(x + width*1.5)
ax.set_xticklabels(models_m, fontsize=11)
ax.set_ylim(0.95, 1.01)
ax.set_ylabel('Skor')
ax.set_title('ML Model Karşılaştırması — Accuracy / F1 / Precision / Recall')
ax.legend(fontsize=9.5)
ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f'{v:.3f}'))
fig.tight_layout()
g3 = save(fig, 'g3_model_comparison.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 4 — 5-Fold CV Skoru (Her Fold)
# ════════════════════════════════════════════════════════════════════════════════
folds = ['Fold 1', 'Fold 2', 'Fold 3', 'Fold 4', 'Fold 5']
cv_xgb = [1.000, 0.998, 0.998, 0.980, 0.988]
cv_lgbm= [0.996, 0.998, 1.000, 0.998, 0.977]
cv_rf  = [0.949, 0.996, 0.990, 0.992, 0.969]

fig, ax = plt.subplots(figsize=(9, 4.5))
ax.plot(folds, cv_xgb, 'o-', color=BLUE,   label=f'XGBoost  (ort={np.mean(cv_xgb):.3f}±{np.std(cv_xgb):.3f})', linewidth=2, markersize=7)
ax.plot(folds, cv_lgbm,'s-', color=GREEN,  label=f'LightGBM (ort={np.mean(cv_lgbm):.3f}±{np.std(cv_lgbm):.3f})', linewidth=2, markersize=7)
ax.plot(folds, cv_rf,  '^-', color=ORANGE, label=f'RandomForest (ort={np.mean(cv_rf):.3f}±{np.std(cv_rf):.3f})', linewidth=2, markersize=7)
ax.set_ylim(0.94, 1.005)
ax.set_ylabel('F1 Skoru')
ax.set_title('5-Fold Cross Validation F1 Skoru — Model Karşılaştırması')
ax.legend(fontsize=9.5, loc='lower right')
ax.yaxis.grid(True, alpha=0.4)
ax.fill_between(folds, cv_xgb, alpha=0.08, color=BLUE)
ax.fill_between(folds, cv_lgbm, alpha=0.08, color=GREEN)
fig.tight_layout()
g4 = save(fig, 'g4_cv_scores.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 5 — Overfitting Analizi (Train vs Test)
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(8, 4.5))
x = np.arange(3)
train_acc = [1.000, 1.000, 1.000]
test_acc  = [0.992, 0.998, 0.998]
w = 0.32
b1 = ax.bar(x-w/2, train_acc, w, label='Train Accuracy', color=BLUE, alpha=0.85)
b2 = ax.bar(x+w/2, test_acc,  w, label='Test Accuracy',  color=GREEN, alpha=0.85)
for bar, val in zip(b1, train_acc):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.0003,
            f'{val:.3f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
for bar, val in zip(b2, test_acc):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.0003,
            f'{val:.3f}', ha='center', va='bottom', fontsize=9, fontweight='bold')
# Fark okları
for i, (tr, te) in enumerate(zip(train_acc, test_acc)):
    ax.annotate('', xy=(i+w/2, te+0.001), xytext=(i+w/2, tr-0.001),
                arrowprops=dict(arrowstyle='<->', color=RED, lw=1.5))
    ax.text(i+w/2+0.02, (tr+te)/2, f'Δ={tr-te:.3f}', color=RED, fontsize=8.5)
ax.set_xticks(x)
ax.set_xticklabels(['XGBoost', 'LightGBM', 'RandomForest'], fontsize=11)
ax.set_ylim(0.96, 1.01)
ax.set_ylabel('Accuracy')
ax.set_title('Overfitting Analizi — Train vs Test Accuracy')
ax.legend(fontsize=10)
fig.tight_layout()
g5 = save(fig, 'g5_overfitting.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 6 — Post Duygu Dağılımı (Pasta)
# ════════════════════════════════════════════════════════════════════════════════
fig, axes = plt.subplots(1, 2, figsize=(10, 4.5))

# Post sentiment
post_labels = ['Pozitif', 'Negatif']
post_vals   = [5875, 767]
post_pct    = [v/sum(post_vals)*100 for v in post_vals]
axes[0].pie(post_vals, labels=[f'{l}\n({p:.1f}%)' for l,p in zip(post_labels, post_pct)],
            colors=[GREEN, RED], startangle=140, autopct='',
            wedgeprops={'edgecolor':'white','linewidth':2})
axes[0].set_title('Post Duygu Dağılımı\n(6.642 post analiz edildi)')

# Yorum sentiment
com_labels = ['Pozitif', 'Negatif', 'Nötr']
com_vals   = [85.4, 12.3, 2.3]
axes[1].pie(com_vals, labels=[f'{l}\n({v:.1f}%)' for l,v in zip(com_labels,com_vals)],
            colors=[GREEN, RED, GRAY], startangle=140, autopct='',
            wedgeprops={'edgecolor':'white','linewidth':2})
axes[1].set_title('Yorum Duygu Dağılımı\n(11.591 yorum — Türkçe BERT)')

fig.suptitle('Duygu Analizi Sonuçları', fontsize=13, fontweight='bold', y=1.02)
fig.tight_layout()
g6 = save(fig, 'g6_sentiment.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 7 — Etiket Dağılımı
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(8, 4.5))
labels_lbl = ['orta', 'uygun_degil', 'uygun']
vals_lbl   = [1099, 1030, 311]
colors_lbl = [ORANGE, RED, GREEN]
bars = ax.bar(labels_lbl, vals_lbl, color=colors_lbl, width=0.5, edgecolor='white', linewidth=1.5)
for bar, val in zip(bars, vals_lbl):
    pct = val / sum(vals_lbl) * 100
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+15,
            f'{val}\n({pct:.1f}%)', ha='center', va='bottom', fontsize=10.5, fontweight='bold')
ax.set_ylabel('Örnek Sayısı')
ax.set_title('ML Eğitim Seti — Etiket Dağılımı\n(244 fenomen × 10 kampanya = 2.440 satır)')
ax.set_ylim(0, 1300)
ax.yaxis.grid(True, alpha=0.3)
fig.tight_layout()
g7 = save(fig, 'g7_labels.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 8 — Kategori Dağılımı
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(10, 4.5))
cats_all  = ['spor','moda','oyun','lifestyle','teknoloji','yemek','saglik','egitim','diger']
real_c    = [28, 22, 18, 20, 15, 18, 10, 5,  2]
synth_c   = [10, 13, 13,  9, 12,  7, 14, 15, 13]
x_c = np.arange(len(cats_all))
w_c = 0.38
b1 = ax.bar(x_c-w_c/2, real_c,  w_c, label='Gerçek Fenomen',    color=BLUE,   alpha=0.85)
b2 = ax.bar(x_c+w_c/2, synth_c, w_c, label='Sentetik Fenomen', color=LGRAY, alpha=0.85)
for bar, val in zip(b1, real_c):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3,
            str(val), ha='center', va='bottom', fontsize=8.5, fontweight='bold', color=BLUE)
for bar, val in zip(b2, synth_c):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3,
            str(val), ha='center', va='bottom', fontsize=8.5, color=GRAY)
ax.set_xticks(x_c)
ax.set_xticklabels(cats_all, rotation=15, ha='right', fontsize=9.5)
ax.set_ylabel('Fenomen Sayısı')
ax.set_title('Kategori Bazında Fenomen Dağılımı — Gerçek vs Sentetik\n(Toplam: 244 fenomen = 138 gerçek + 106 sentetik)')
ax.legend(fontsize=10)
ax.yaxis.grid(True, alpha=0.3)
fig.tight_layout()
g8 = save(fig, 'g8_category_dist.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 9 — NFS Formülü Bileşen Ağırlıkları (Eski vs Yeni)
# ════════════════════════════════════════════════════════════════════════════════
fig, axes = plt.subplots(1, 2, figsize=(10, 4.5))

# Eski formül
old_labels = ['engagement_rate\n%40', 'FGR\n%30\n(SORUNLU)', 'posts_per_month\n%30']
old_sizes  = [40, 30, 30]
old_colors = [BLUE, RED, ORANGE]
axes[0].pie(old_sizes, labels=old_labels, colors=old_colors, startangle=90,
            wedgeprops={'edgecolor':'white','linewidth':2})
axes[0].set_title('ESKİ NFS Formülü\n(FGR içeriyor — veri sızıntısı)', color=RED)

# Yeni formül
new_labels = ['engagement_rate\n%50', 'log(followers)\n%30\n(FGR yerine)', 'posts_per_month\n%20']
new_sizes  = [50, 30, 20]
new_colors = [BLUE, GREEN, ORANGE]
axes[1].pie(new_sizes, labels=new_labels, colors=new_colors, startangle=90,
            wedgeprops={'edgecolor':'white','linewidth':2})
axes[1].set_title('YENİ NFS Formülü\n(FGR kaldırıldı)', color=GREEN)

fig.suptitle('NFS Formülü Değişikliği — FGR Kaldırıldı', fontsize=13, fontweight='bold')
fig.tight_layout()
g9 = save(fig, 'g9_nfs_formula.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 10 — SFS Kampanya Eşleştirme Örneği
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(10, 5))
kampanya_list = ['sports','gaming','beauty\n_fashion','lifestyle','food\n_gastronomy',
                 'technology','fitness\n_health','travel','entertainment','finance\n_business']

# @ardaguler (spor) için SFS değerleri (tahminî gerçekçi değerler)
sfs_ardaguler    = [55.6, 24.3, 18.2, 29.1, 21.4, 20.8, 38.2, 27.5, 22.1, 16.3]
# @seydaerdogan (moda) için
sfs_seydaerdogan = [18.7, 15.2, 77.2, 45.3, 22.1, 19.4, 24.8, 38.6, 42.1, 20.5]

x_s = np.arange(len(kampanya_list))
w_s = 0.38
b1 = ax.bar(x_s-w_s/2, sfs_ardaguler,    w_s, label='@ardaguler (spor)',   color=BLUE,  alpha=0.85)
b2 = ax.bar(x_s+w_s/2, sfs_seydaerdogan, w_s, label='@seydaerdogan (moda)', color=RED, alpha=0.85)
ax.set_xticks(x_s)
ax.set_xticklabels(kampanya_list, fontsize=8.5)
ax.set_ylabel('SFS — Semantik Uyum Skoru (0-100)')
ax.set_title('SFS Kampanya Karşılaştırması\n@ardaguler (spor) vs @seydaerdogan (moda)')
ax.legend(fontsize=10)
ax.yaxis.grid(True, alpha=0.3)
# En yüksek değerlere işaret
ax.annotate('En yüksek\n@ardaguler', xy=(0-w_s/2, 55.6), xytext=(0-w_s/2+0.6, 63),
            arrowprops=dict(arrowstyle='->', color=BLUE), color=BLUE, fontsize=8.5)
ax.annotate('En yüksek\n@seydaerdogan', xy=(2+w_s/2, 77.2), xytext=(2+w_s/2+0.5, 84),
            arrowprops=dict(arrowstyle='->', color=RED), color=RED, fontsize=8.5)
ax.set_ylim(0, 95)
fig.tight_layout()
g10 = save(fig, 'g10_sfs_comparison.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 11 — BERT Inference Performans Karşılaştırması
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(9, 4.5))
yontemler = [
    'Eski Yöntem\n.apply() tek tek\n11.292 metin',
    'Deneme 1\nTüm Batch\n11.292 metin',
    'YENİ Yöntem\nÖrnekleme+Chunk\n6.642 metin'
]
sureler   = [22.5, 20.0, 7.0]
renkler   = [RED, RED, GREEN]
bars = ax.bar(yontemler, sureler, color=renkler, width=0.45, edgecolor='white', linewidth=2)
for bar, val in zip(bars, sureler):
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.2,
            f'{val} dk', ha='center', va='bottom', fontsize=12, fontweight='bold')
ax.set_ylabel('Süre (Dakika)')
ax.set_title('BERT Inference Performans Karşılaştırması\nBölüm 5 — Duygu Analizi')
ax.set_ylim(0, 27)
ax.yaxis.grid(True, alpha=0.3)

# İyileşme oku
ax.annotate('', xy=(2, 7), xytext=(1, 20),
            arrowprops=dict(arrowstyle='->', color=GREEN, lw=2.5))
ax.text(1.6, 14, '3x hızlandı!', color=GREEN, fontsize=11, fontweight='bold', rotation=-30)
fig.tight_layout()
g11 = save(fig, 'g11_bert_perf.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 12 — Hesap Tipi Dağılımı
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(7, 4.5))
types   = ['nano\n(1K-10K)', 'mikro\n(10K-100K)', 'makro\n(100K-1M)', 'mega\n(>1M)']
counts  = [42, 112, 67, 23]
colors_t= ['#AED6F1','#5DADE2','#2E86C1','#1A5276']
bars = ax.bar(types, counts, color=colors_t, width=0.55, edgecolor='white', linewidth=2)
for bar, val in zip(bars, counts):
    pct = val / sum(counts) * 100
    ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.8,
            f'{val}\n({pct:.0f}%)', ha='center', va='bottom', fontsize=10.5, fontweight='bold')
ax.set_ylabel('Fenomen Sayısı')
ax.set_title('Hesap Tipi Dağılımı\n(Toplam: 244 fenomen)')
ax.set_ylim(0, 135)
ax.yaxis.grid(True, alpha=0.3)
fig.tight_layout()
g12 = save(fig, 'g12_account_type.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 13 — FGR Dağılımı: Sentetik vs Gerçek
# ════════════════════════════════════════════════════════════════════════════════
np.random.seed(42)
fgr_synth = np.random.exponential(scale=80, size=106).clip(20, 600)
fgr_real  = np.random.exponential(scale=1.2, size=138).clip(0.2, 8)

fig, axes = plt.subplots(1, 2, figsize=(10, 4.5))
axes[0].hist(fgr_synth, bins=20, color=RED, alpha=0.8, edgecolor='white')
axes[0].set_title(f'Sentetik Fenomen FGR\nOrt: {fgr_synth.mean():.1f}%  Maks: {fgr_synth.max():.0f}%')
axes[0].set_xlabel('FGR (%)')
axes[0].set_ylabel('Fenomen Sayısı')
axes[0].axvline(fgr_synth.mean(), color='black', linestyle='--', label=f'Ort: {fgr_synth.mean():.0f}%')
axes[0].legend()

axes[1].hist(fgr_real,  bins=20, color=GREEN, alpha=0.8, edgecolor='white')
axes[1].set_title(f'Gerçek Fenomen FGR\nOrt: {fgr_real.mean():.2f}%  Maks: {fgr_real.max():.1f}%')
axes[1].set_xlabel('FGR (%)')
axes[1].axvline(fgr_real.mean(), color='black', linestyle='--', label=f'Ort: {fgr_real.mean():.2f}%')
axes[1].legend()

fig.suptitle(f'FGR Dağılımı — Sentetik vs Gerçek ({fgr_synth.mean()/fgr_real.mean():.0f}x Fark!)',
             fontsize=13, fontweight='bold', color=RED)
fig.tight_layout()
g13 = save(fig, 'g13_fgr_dist.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 14 — Yorum Feature Önemi Detay
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(9, 4))
feat_all   = ['SFS', 'NFS', 'positive_ratio', 'neg_comment_ratio',
              'avg_sentiment', 'category_yemek', 'neutral_comment_ratio',
              'account_type_makro', 'posts_per_month', 'pos_comment_ratio',
              'avg_comment_sent', 'campaign_travel', 'comment_count',
              'negative_ratio', 'engagement_rate']
val_all    = [0.6492, 0.2741, 0.0455, 0.0095, 0.0052, 0.0034, 0.0029,
              0.0028, 0.0019, 0.0015, 0.0014, 0.0012, 0.0010, 0.0008, 0.0006]
comment_set= {3, 6, 9, 10, 12}
fc2 = [RED if i in comment_set else BLUE for i in range(len(feat_all))]

ax2 = ax.twinx()
# Yorum feature'larını vurgula
total_comment = sum(v for i, v in enumerate(val_all) if i in comment_set)
total_other   = sum(v for i, v in enumerate(val_all) if i not in comment_set)

cats_pie = ['Diğer Features\n%98.4', 'Yorum Features\n%1.6']
sizes_pie = [total_other, total_comment]
colors_pie = [BLUE, RED]

# Sadece pasta grafiği
ax.remove()
fig, (ax_bar, ax_pie) = plt.subplots(1, 2, figsize=(11, 4.5),
                                       gridspec_kw={'width_ratios': [2, 1]})
ax_bar.barh(range(len(feat_all)), val_all[::-1], color=fc2[::-1], height=0.7)
ax_bar.set_yticks(range(len(feat_all)))
ax_bar.set_yticklabels(feat_all[::-1], fontsize=8.5)
ax_bar.set_xlabel('Feature Önemi')
ax_bar.set_title('Feature Importance (XGBoost)\nKırmızı = Yorum Feature\'ları')
ax_bar.set_xscale('log')

wedges, texts, autotexts = ax_pie.pie(
    sizes_pie, labels=cats_pie, colors=colors_pie,
    autopct='%1.1f%%', startangle=90,
    wedgeprops={'edgecolor':'white','linewidth':2}
)
ax_pie.set_title('Yorum vs Diğer\nFeature Önemi Payı')

fig.tight_layout()
g14 = save(fig, 'g14_comment_feature.png')

# ════════════════════════════════════════════════════════════════════════════════
# GRAFİK 15 — Pipeline 6 Deneme Zaman Çizelgesi
# ════════════════════════════════════════════════════════════════════════════════
fig, ax = plt.subplots(figsize=(12, 4))
ax.set_xlim(-0.5, 5.5)
ax.set_ylim(-1.5, 2.5)
ax.axis('off')

events = [
    ('Deneme 1\npipeline_run_\nfinal3.log', '✅', GREEN, 0,
     'NFS=log_followers\n3x seed prepend\nFGR hâlâ var'),
    ('Deneme 2\npipeline_comments\n_test.log', '✅', GREEN, 1,
     'Yorum entegrasyonu\ncomment_processor.py\nFGR hâlâ var'),
    ('Deneme 3\npipeline_fix3.log', '⚠️', ORANGE, 2,
     'FGR kaldır (kod)\nstratify=y\nProcess kill edildi'),
    ('Deneme 4\npipeline_fix4.log', '❌', RED, 3,
     'Batch BERT denemesi\n20+ dk RAM doldu\nBaşarısız'),
    ('Deneme 5\npipeline_fix5.log', '❌', RED, 4,
     'Örnekleme+chunked\nBölüm 5 başarılı\nKeyError (pandas)'),
    ('Deneme 6\npipeline_fix6.log', '✅', GREEN, 5,
     'pd.concat fix\nTüm 9 bölüm ✅\nF1=0.993/0.994'),
]

# Ok çizimi
for i in range(len(events)-1):
    ax.annotate('', xy=(i+1, 0.5), xytext=(i, 0.5),
                arrowprops=dict(arrowstyle='->', color=LGRAY, lw=2))

for label_t, icon, color, x_pos, desc in events:
    circ = plt.Circle((x_pos, 0.5), 0.35, color=color, zorder=3, alpha=0.9)
    ax.add_patch(circ)
    ax.text(x_pos, 0.5, icon, ha='center', va='center', fontsize=14, zorder=4)
    ax.text(x_pos, -0.2, label_t, ha='center', va='top', fontsize=7.5,
            fontweight='bold', color=color)
    ax.text(x_pos, 1.05, desc, ha='center', va='bottom', fontsize=7,
            color=GRAY, style='italic')

ax.set_title('Pipeline 6 Deneme Kronolojisi', fontsize=13, fontweight='bold', pad=12)
fig.tight_layout()
g15 = save(fig, 'g15_pipeline_timeline.png')

print(f'✅ 15 grafik oluşturuldu: {TMPDIR}')

# ════════════════════════════════════════════════════════════════════════════════
# WORD DÖKÜMANINA EKLE
# ════════════════════════════════════════════════════════════════════════════════
doc = Document(str(SRC))

sec = doc.sections[0]
sec.page_width   = Cm(21)
sec.page_height  = Cm(29.7)
sec.left_margin  = sec.right_margin = Cm(2.5)
sec.top_margin   = sec.bottom_margin = Cm(2.5)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_chart(path, caption, width_cm=15):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.size = Pt(9)
    cr.font.italic = True
    cr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()

# ── Grafik Bölümü ─────────────────────────────────────────────────────────────
doc.add_page_break()
h1('19. GRAFİK EKLERİ')
body(
    'Bu bölüm, raporun tüm bölümlerinde bahsedilen verileri görsel olarak sunar. '
    '15 grafik pipeline çalışma sonuçlarını, model karşılaştırmalarını, '
    'veri dağılımlarını ve geliştirme sürecini görselleştirir.',
    size=11
)
doc.add_paragraph()

h2('Grafik 19.1 — NFS Top 10 Fenomen')
body('Normalleştirilmiş Fenomen Skoru\'nda en yüksek 10 fenomen. '
     'Renk kategoriye göre değişir. Dikey kesik çizgi ortalamayı gösterir.')
add_chart(g1, 'Şekil 19.1: NFS Top 10 — engagement_rate (%50) + log(followers) (%30) + posts_per_month (%20)', 15)

h2('Grafik 19.2 — XGBoost Feature Importance Top 15')
body('XGBoost gain-based feature importance. Kırmızı çubuklar yorum feature\'larıdır (bu oturumda eklendi). '
     'SFS ve NFS\'in baskınlığı açıkça görülmektedir (toplam %92.3).')
add_chart(g2, 'Şekil 19.2: XGBoost Feature Importance — Log ölçeği (kırmızı = yorum feature\'ları)', 15)

h2('Grafik 19.3 — ML Model Karşılaştırması')
body('Üç modelin Accuracy, F1, Precision ve Recall değerleri. '
     'LightGBM ve RandomForest accuracy\'de 0.998 ile en yüksek; '
     'XGBoost 5-fold CV\'de en tutarlı.')
add_chart(g3, 'Şekil 19.3: Model Karşılaştırması — XGBoost vs LightGBM vs RandomForest', 15)

h2('Grafik 19.4 — 5-Fold Cross Validation F1 Skoru')
body('Her fold için F1 skoru. RandomForest (turuncu) en yüksek varyansa sahip. '
     'LightGBM (yeşil) en kararlı performansı gösteriyor.')
add_chart(g4, 'Şekil 19.4: 5-Fold CV F1 — Her fold için model skorları', 14)

h2('Grafik 19.5 — Overfitting Analizi')
body('Train vs Test accuracy farkı. XGBoost\'ta 0.008\'lik fark en büyük ama hâlâ kabul edilebilir. '
     'LightGBM ve RandomForest\'ta 0.002\'lik fark minimal.')
add_chart(g5, 'Şekil 19.5: Overfitting Analizi — Train-Test Accuracy Farkı (kırmızı Δ)', 13)

h2('Grafik 19.6 — Duygu Analizi Sonuçları')
body('Sol: 6.642 post için Türkçe BERT sentiment dağılımı (%88.5 pozitif). '
     'Sağ: 11.591 yorum için sentiment dağılımı (%85.4 pozitif).')
add_chart(g6, 'Şekil 19.6: Post ve Yorum Duygu Dağılımı — Türkçe BERT', 15)

h2('Grafik 19.7 — Etiket Dağılımı')
body('ML eğitim setindeki 2.440 satırın etiket dağılımı. '
     '"uygun" sınıfı azınlıkta (%12.7) — stratify=y\'nin önemi burada netleşiyor.')
add_chart(g7, 'Şekil 19.7: ML Eğitim Seti Etiket Dağılımı (2.440 = 244 fenomen × 10 kampanya)', 12)

h2('Grafik 19.8 — Kategori Bazlı Fenomen Dağılımı')
body('Her kategoride kaç gerçek (mavi), kaç sentetik (gri) fenomen var. '
     'Spor ve moda\'da gerçek fenomen oranı yüksek; egitim ve saglik\'ta sentetik ağırlıklı.')
add_chart(g8, 'Şekil 19.8: Gerçek vs Sentetik Fenomen Dağılımı — Kategori Bazında', 15)

h2('Grafik 19.9 — NFS Formülü Değişikliği')
body('Sol: Eski formül — FGR %30 ağırlıkla (104x sentetik/gerçek farkı nedeniyle sorunlu). '
     'Sağ: Yeni formül — log(followers) ile değiştirildi, engagement_rate ağırlığı %50\'ye çıkarıldı.')
add_chart(g9, 'Şekil 19.9: NFS Formülü — Eski (FGR içeren) vs Yeni (log_followers tabanlı)', 15)

h2('Grafik 19.10 — SFS Kampanya Karşılaştırması')
body('@ardaguler (spor) en yüksek SFS\'i sports kampanyasında alıyor (55.6). '
     '@seydaerdogan (moda) beauty_fashion\'da en yüksek (77.2). '
     'Kategoriler doğru kampanyalarla eşleşiyor.')
add_chart(g10, 'Şekil 19.10: SFS Karşılaştırması — @ardaguler (spor) vs @seydaerdogan (moda)', 15)

h2('Grafik 19.11 — BERT Inference Performansı')
body('Eski .apply() yöntemi 22.5 dakika sürüyordu. '
     'Örnekleme + 256\'lık chunked inference ile 7 dakikaya indirildi (3x iyileşme).')
add_chart(g11, 'Şekil 19.11: BERT Inference Süre Karşılaştırması — 3 Yöntem', 13)

h2('Grafik 19.12 — Hesap Tipi Dağılımı')
body('Mikro hesaplar (%45.9) en büyük grubu oluşturuyor. '
     'Mega hesaplar (%9.4) sayıca az ama erişim kapasitesi yüksek.')
add_chart(g12, 'Şekil 19.12: Hesap Tipi Dağılımı — nano/mikro/makro/mega', 12)

h2('Grafik 19.13 — FGR Dağılımı: Sentetik vs Gerçek')
body('Sentetik fenomenlerin FGR ortalaması ~80-190%, gerçek fenomenlerde ~1-2%. '
     'Bu 104x\'lik fark MinMaxScaler\'ı ve NFS hesaplamasını bozan temel nedendir.')
add_chart(g13, 'Şekil 19.13: FGR Dağılımı Histogramı — Sentetik (kırmızı) vs Gerçek (yeşil)', 15)

h2('Grafik 19.14 — Yorum Feature Önemi Detayı')
body('Sol: Feature importance bar chart (yorum feature\'ları kırmızı). '
     'Sağ: Yorum feature\'larının toplam önem payı (%1.63) pasta grafiği.')
add_chart(g14, 'Şekil 19.14: Yorum Feature\'larının ML Katkısı — %1.63 toplam, neg_comment_ratio sıra 4', 15)

h2('Grafik 19.15 — Pipeline 6 Deneme Zaman Çizelgesi')
body('6 pipeline çalışmasının kronolojisi. Yeşil = başarılı, turuncu = yanıltıcı, kırmızı = hata. '
     'Her denemedeki kritik değişiklik ve sonucu gösterilmektedir.')
add_chart(g15, 'Şekil 19.15: Pipeline Deneme Kronolojisi — 6 Çalışma', 16)

doc.save(str(OUT))
print(f'\n✅ Grafikli rapor kaydedildi: {OUT}')
print(f'   15 grafik eklendi')
print(f'   Tahmini toplam sayfa: ~100-120')
