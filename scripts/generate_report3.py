# -*- coding: utf-8 -*-
"""
Fenomen-Marka Eşleştirme Sistemi
Kapsamlı Geliştirme Süreci Raporu — Bölüm 1/4 (Kapak → Bölüm 9)
Hedef: ~80-100 sayfa
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

sec = doc.sections[0]
sec.page_width   = Cm(21)
sec.page_height  = Cm(29.7)
sec.left_margin  = sec.right_margin = Cm(2.5)
sec.top_margin   = sec.bottom_margin = Cm(2.5)

# ── Yardımcılar ──────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h4(text):
    p = doc.add_heading(text, level=4)
    return p

def body(text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(text, style=style)
    p.runs[0].font.size = Pt(10.5)
    return p

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    p.runs[0].font.size = Pt(10.5)
    return p

def _shade(p, fill):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  fill)
    p._p.get_or_add_pPr().append(shading)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(0.8)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    try: _shade(p, 'EFEFEF')
    except: pass
    return p

def log_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x00)
    try: _shade(p, 'E8F5E9')
    except: pass
    return p

def err_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    try: _shade(p, 'FFEBEE')
    except: pass
    return p

def warn_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x50, 0x00)
    try: _shade(p, 'FFF8E1')
    except: pass
    return p

def label(text, color):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(*color)
    return p

def _shade_cell(cell, fill):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  fill)
    cell._tc.get_or_add_tcPr().append(shading)

def new_tbl(headers, col_fills=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    row = t.rows[0]
    for i, h in enumerate(headers):
        row.cells[i].text = h
        run = row.cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        _shade_cell(row.cells[i], 'D6E4F0')
    return t

def tbl_row(tbl, vals, fill=None):
    row = tbl.add_row()
    for i, v in enumerate(vals):
        row.cells[i].text = str(v)
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
        if fill:
            _shade_cell(row.cells[i], fill)
    return row

def divider():
    body('─' * 90, color=(0xCC, 0xCC, 0xCC))

# ════════════════════════════════════════════════════════════════════════════════
# KAPAK
# ════════════════════════════════════════════════════════════════════════════════
doc.add_paragraph('\n\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FENOMEN-MARKA EŞLEŞTİRME SİSTEMİ')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Influencer–Brand Matching System')
r.font.size = Pt(16)
r.italic = True
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Geliştirme Süreci, Karşılaşılan Sorunlar, Çözümler\nve Teknik Dokümantasyon')
r.font.size = Pt(13)
r.italic = True

doc.add_paragraph('\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TÜBİTAK 2209-A Lisans Bitirme Projesi')
r.font.size = Pt(12)
r.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Tarih: {datetime.datetime.now().strftime("%d %B %Y")}\n'
              'Repo: github.com/eceekkutlu/FenomenMarkaUyumu')
r.font.size = Pt(10)
r.italic = True

doc.add_paragraph('\n\n')
t_cover = doc.add_table(rows=1, cols=2)
t_cover.style = 'Table Grid'
c0, c1 = t_cover.rows[0].cells
c0.text = 'Python'
c1.text = '3.14'
r2 = t_cover.add_row()
r2.cells[0].text = 'Flask'
r2.cells[1].text = '3.x REST API'
r3 = t_cover.add_row()
r3.cells[0].text = 'Modeller'
r3.cells[1].text = 'XGBoost | LightGBM | RandomForest'
r4 = t_cover.add_row()
r4.cells[0].text = 'NLP'
r4.cells[1].text = 'SBERT paraphrase-multilingual | Türkçe BERT'
r5 = t_cover.add_row()
r5.cells[0].text = 'Veri'
r5.cells[1].text = '244 fenomen | 11.292 post | 11.591 yorum'

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# İÇİNDEKİLER
# ════════════════════════════════════════════════════════════════════════════════
h1('İÇİNDEKİLER')
bolumler = [
    ('1',  'Yönetici Özeti'),
    ('2',  'Projenin Motivasyonu ve Araştırma Soruları'),
    ('3',  'Sistem Mimarisi ve Bileşen Haritası'),
    ('4',  'Geliştirme Ortamı ve Bağımlılıklar'),
    ('5',  'Veri Seti — Detaylı İnceleme'),
    ('6',  'Teorik Arka Plan (SBERT, BERT, XGBoost, LightGBM, K-Means)'),
    ('7',  'Gerçek Influencer Verisi Ekleme Süreci'),
    ('8',  'comment_processor.py — Tam Modül Dokümantasyonu'),
    ('9',  'Analiz Pipeline Mimarisi — 9 Bölüm Detayı'),
    ('10', 'Puanlama Formülleri — Matematiksel Türetim ve Örnekler'),
    ('11', 'Flask REST API — Tam Dokümantasyon'),
    ('12', 'Karşılaşılan Sorunlar ve Çözümler (9 Kritik Sorun)'),
    ('13', 'Pipeline Çalışma Geçmişi — 6 Deneme Kronolojisi'),
    ('14', 'ML Model Eğitimi, Doğrulama ve Sonuçlar'),
    ('15', 'Feature Importance ve Yorum Verisi Katkısı'),
    ('16', 'Frontend Test Sonuçları — 3 Kampanya'),
    ('17', 'Proje Dizini Temizliği ve GitHub Push'),
    ('18', 'Öğrenilen Dersler ve Geliştirme Önerileri'),
]
for no, baslik in bolumler:
    bullet(f'{no}. {baslik}')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 1. YÖNETİCİ ÖZETİ
# ════════════════════════════════════════════════════════════════════════════════
h1('1. YÖNETİCİ ÖZETİ')

body(
    'Bu rapor, TÜBİTAK 2209-A Lisans Bitirme Projesi kapsamında geliştirilen '
    '"Fenomen-Marka Eşleştirme Sistemi"nin tüm geliştirme sürecini, '
    'karşılaşılan teknik sorunları, uygulanan çözümleri ve sistem mimarisini '
    'kapsamlı biçimde belgelemektedir. Rapor, yalnızca sonuçları değil; '
    'her hatanın nasıl tespit edildiğini, hangi adımların denendiğini ve '
    'nihai çözüme nasıl ulaşıldığını da aktarmaktadır.'
)

h2('1.1 Proje Özeti')
t = new_tbl(['Konu', 'Değer'])
tbl_row(t, ['Proje Adı',          'Fenomen-Marka Eşleştirme Sistemi (Influencer-Brand Matching)'])
tbl_row(t, ['Fon Kaynağı',        'TÜBİTAK 2209-A Lisans Bitirme Projesi'])
tbl_row(t, ['Repo',               'github.com/eceekkutlu/FenomenMarkaUyumu'])
tbl_row(t, ['Dil/Çerçeve',        'Python 3.14 + Flask + React'])
tbl_row(t, ['NLP Modelleri',      'SBERT (paraphrase-multilingual) + BERT (Türkçe sentiment)'])
tbl_row(t, ['ML Modelleri',       'XGBoost (F1=0.993) + LightGBM (F1=0.994) + RandomForest'])
tbl_row(t, ['Veri Boyutu',        '244 fenomen | 11.292 post | 11.591 yorum'])
tbl_row(t, ['Pipeline Bölümleri', '9 (Veri Yükleme → Checkpoint)'])
tbl_row(t, ['API Endpoint',       '6 (/recommend, /influencers, /stats, /campaigns vb.)'])
tbl_row(t, ['Raporlama Tarihi',   datetime.datetime.now().strftime('%d.%m.%Y')])
doc.add_paragraph()

h2('1.2 Kritik Başarılar')
bullet('FGR veri sızıntısı tespit edildi ve ML feature setinden kaldırıldı (37→36 feature)')
bullet('BERT inference süresi 20+ dakikadan ~7 dakikaya indirildi (örnekleme + chunked)')
bullet('Python 3.14 + pandas groupby uyumsuzluğu explicit for-loop ile çözüldü')
bullet('3x kategori seed prepend ile SBERT semantik eşleşme doğruluğu artırıldı')
bullet('comment_processor.py modülü sıfırdan yazılarak yorum verisi sisteme entegre edildi')
bullet('Gerçek 138 Türk fenomeni eklendi, toplam 244 fenomene ulaşıldı')
bullet('3 kampanya türüyle frontend testi: tüm öneriler doğru kategori')

h2('1.3 Sistem Çalışma Kanıtı')
body('En son başarılı pipeline çalışması (pipeline_fix6.log) sonuçları:')
log_block(
    'BÖLÜM 1  ✅  244 fenomen, 11292 post, 11591 yorum yüklendi\n'
    'BÖLÜM 2  ✅  influencer_summary: 244 fenomen, 23 sütun\n'
    'BÖLÜM 3  ✅  NFS hesaplandı (log_followers tabanlı, FGR YOK)\n'
    'BÖLÜM 4  ✅  SFS hesaplandı (SBERT + yorum %30 ağırlık)\n'
    'BÖLÜM 5  ✅  6642 post + yorum duygu analizi tamamlandı (~7 dakika)\n'
    'BÖLÜM 6  ✅  Sahte takipçi tespiti + K-Means clustering\n'
    'BÖLÜM 7  ✅  XGBoost F1=0.993 | LightGBM F1=0.994 [FGR kaldırıldı]\n'
    'BÖLÜM 8  ✅  Model raporları kaydedildi (confusion_matrices.png, vb.)\n'
    'BÖLÜM 9  ✅  Checkpoint güncellendi (244 fenomen)\n'
    '─────────────────────────────────────────────────────────\n'
    'Pipeline tamamlandı. Feature sayısı: 36  |  FGR var mı: False'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 2. MOTİVASYON
# ════════════════════════════════════════════════════════════════════════════════
h1('2. PROJENİN MOTİVASYONU VE ARAŞTIRMA SORULARI')

body(
    'Sosyal medya influencer pazarlaması, 2024 yılı itibarıyla küresel ölçekte '
    '21 milyar dolarlık bir endüstri haline gelmiştir. Türkiye özelinde ise '
    'Instagram, TikTok ve YouTube gibi platformlarda faaliyet gösteren binlerce '
    'içerik üreticisi bulunmaktadır. Ancak markalar için doğru fenomeni seçmek '
    'hâlâ büyük ölçüde manuel, sezgisel ve zaman alıcı bir süreçtir.'
)
body(
    'Mevcut araçların temel eksikliği şudur: Bir marka "spor ekipmanı satıyoruz" '
    'dediğinde, sistem yalnızca anahtar kelime eşleştirmesi yapar. Ancak bir '
    'spor fenomeninin gerçekten spor içeriği üretip üretmediğini, takipçileriyle '
    'ne kadar etkileşime girdiğini veya yorumlarının ne ölçüde pozitif olduğunu '
    'ölçmek için semantik anlama, duygu analizi ve makine öğrenmesi gerekir.'
)

h2('2.1 Araştırma Soruları')
t = new_tbl(['No', 'Araştırma Sorusu', 'Sistem Yanıtı'])
tbl_row(t, ['1', 'Bir markanın doğal dil açıklaması, semantik olarak en uygun fenomene nasıl eşleştirilir?', 'SBERT + cosine similarity (SFS)'])
tbl_row(t, ['2', 'Bir fenomenin genel kalitesi nasıl sayısal olarak ölçülür?', 'NFS: engagement + followers + aktivite'])
tbl_row(t, ['3', 'Sahte takipçi riski sistematik olarak nasıl tespit edilir?', 'K-Means clustering + risk skoru'])
tbl_row(t, ['4', 'Yorum verisi fenomen kalitesi değerlendirmesini iyileştirir mi?', 'BERT yorum sentiment + ML feature'])
tbl_row(t, ['5', '"Uygun/orta/uygun değil" sınıflandırması ML ile ne kadar doğru yapılabilir?', 'XGBoost F1=0.993, LightGBM F1=0.994'])
doc.add_paragraph()

h2('2.2 Neden Bu Yaklaşım?')
body(
    'Kural tabanlı sistemler (if followers > 1M: "makro") çok kaba kalır. '
    'Makine öğrenmesi tek başına yorumlanamaz. Bu projede hibrit bir yaklaşım '
    'tercih edildi: Önce NLP ile ham skoru hesapla (SFS, NFS), ardından ML ile '
    'bunu onaya sun (uygun/orta/uygun_degil), son olarak da kullanıcıya '
    'açıklanabilir bir rapor sun (skor bileşenleri, risk etiketi).'
)
body(
    'Türkçe dil desteği kritik bir gereksinimdi. Bu nedenle İngilizce kısıtı '
    'olmayan çok dilli SBERT modeli (paraphrase-multilingual-MiniLM-L12-v2) ve '
    'Türkçe özelinde eğitilmiş BERT (savasy/bert-base-turkish-sentiment-cased) '
    'seçildi.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 3. SİSTEM MİMARİSİ
# ════════════════════════════════════════════════════════════════════════════════
h1('3. SİSTEM MİMARİSİ VE BİLEŞEN HARİTASI')

body(
    'Sistem dört katmanlı bir akışa sahiptir. Her katman bir sonrakini besler '
    've bu tasarım sayesinde her bileşen bağımsız olarak test edilebilir, '
    'değiştirilebilir veya optimize edilebilir.'
)

h2('3.1 Mimari Akış Diyagramı')
code(
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│                    KATMAN 1: VERİ TOPLAMA                       │\n'
    '│  influencer_profiles.csv  +  influencer_posts.csv               │\n'
    '│  influencer_comments.csv  (YENİ — bu oturumda eklendi)          │\n'
    '└───────────────────────────┬─────────────────────────────────────┘\n'
    '                            │\n'
    '                            ▼\n'
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│              KATMAN 2: analiz_pipeline.py (9 BÖLÜM)             │\n'
    '│  Bölüm 1: Yükleme  →  Bölüm 2: Özellik Müh.                    │\n'
    '│  Bölüm 3: NFS      →  Bölüm 4: SFS (SBERT)                     │\n'
    '│  Bölüm 5: Sentiment (BERT)  →  Bölüm 6: Sahte/Cluster          │\n'
    '│  Bölüm 7: ML Eğitim  →  Bölüm 8: Doğrulama  →  Bölüm 9: Kayıt │\n'
    '│                                                                  │\n'
    '│  ÇIKTI: influencer_summary_checkpoint.pkl                        │\n'
    '│          best_model_xgb.pkl  best_model_lgbm.pkl                │\n'
    '│          feature_columns.pkl  label_encoder.pkl                 │\n'
    '└───────────────────────────┬─────────────────────────────────────┘\n'
    '                            │\n'
    '                            ▼\n'
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│                 KATMAN 3: Flask REST API (app.py)                │\n'
    '│  POST /recommend    GET /influencers    GET /stats               │\n'
    '│  GET /campaigns     GET /influencers/<name>/similar             │\n'
    '│  Startup: checkpoint.pkl yükle → SBERT encode → serve           │\n'
    '└───────────────────────────┬─────────────────────────────────────┘\n'
    '                            │\n'
    '                            ▼\n'
    '┌─────────────────────────────────────────────────────────────────┐\n'
    '│              KATMAN 4: Frontend (frontend/index.html)            │\n'
    '│  Marka açıklaması gir → Önerilen fenomenler görüntüle           │\n'
    '│  Skor bileşenleri + risk etiketi + ML sınıf                     │\n'
    '└─────────────────────────────────────────────────────────────────┘'
)

h2('3.2 Dosya Sistemi Bileşen Haritası')
code(
    'TezBitirme/\n'
    '│\n'
    '├── analiz_pipeline.py       # Ana pipeline (~50KB, 700+ satır)\n'
    '│     ├── BÖLÜM 1: Veri yükleme + comment_processor çağrısı\n'
    '│     ├── BÖLÜM 2: influencer_summary oluşturma\n'
    '│     ├── BÖLÜM 3: NFS hesaplama (MinMaxScaler)\n'
    '│     ├── BÖLÜM 4: SBERT embeddings + SFS\n'
    '│     ├── BÖLÜM 5: BERT sentiment (post + yorum)\n'
    '│     ├── BÖLÜM 6: influencer_features.py çağrısı (cluster+risk)\n'
    '│     ├── BÖLÜM 7: ML eğitim (XGBoost, LightGBM, RandomForest)\n'
    '│     ├── BÖLÜM 8: Model doğrulama + raporlama\n'
    '│     └── BÖLÜM 9: Checkpoint kayıt\n'
    '│\n'
    '├── app.py                   # Flask REST API (~23KB, 500+ satır)\n'
    '│     ├── Startup: checkpoint + model yükleme\n'
    '│     ├── /recommend: SBERT real-time encode + ML sınıflandırma\n'
    '│     ├── /influencers: Filtrelenmiş liste\n'
    '│     ├── /influencers/<name>/similar: K-Means cluster sorgusu\n'
    '│     ├── /stats: Sistem istatistikleri\n'
    '│     └── /campaigns: 10 kampanya bilgisi\n'
    '│\n'
    '├── comment_processor.py     # Yorum işleme modülü (~17KB) — YENİ\n'
    '│     ├── load_comments()            # Esnek CSV yükleme\n'
    '│     ├── clean_comment()            # Metin temizleme\n'
    '│     ├── sample_top_comments()      # Beğeni tabanlı örnekleme\n'
    '│     ├── aggregate_comments()       # Fenomen bazı birleştirme\n'
    '│     ├── analyze_comment_sentiment() # BERT duygu analizi\n'
    '│     └── build_enriched_influencer_text() # SBERT zenginleştirme\n'
    '│\n'
    '├── influencer_features.py   # Sahte takipçi + clustering\n'
    '│     ├── compute_fake_risk()\n'
    '│     ├── run_kmeans_clustering()\n'
    '│     └── compute_secondary_tags()\n'
    '│\n'
    '├── [VERİ]\n'
    '│   ├── influencer_profiles.csv  (244 fenomen)\n'
    '│   ├── influencer_posts.csv     (11.292 post)\n'
    '│   └── influencer_comments.csv  (11.591 yorum) — YENİ\n'
    '│\n'
    '├── [MODELLER — Pipeline çıktıları]\n'
    '│   ├── best_model_xgb.pkl   (~310 KB)\n'
    '│   ├── best_model_lgbm.pkl  (~963 KB)\n'
    '│   ├── feature_columns.pkl  (36 feature ismi)\n'
    '│   ├── label_encoder.pkl    (orta/uygun/uygun_degil → 0/1/2)\n'
    '│   └── influencer_summary_checkpoint.pkl  (244 fenomen, tüm skorlar)\n'
    '│\n'
    '├── model_reports/\n'
    '│   ├── confusion_matrices.png\n'
    '│   ├── feature_importance.png\n'
    '│   ├── model_comparison.png\n'
    '│   └── model_validation_report.txt\n'
    '│\n'
    '└── frontend/\n'
    '    ├── index.html\n'
    '    └── static/'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 4. GELİŞTİRME ORTAMI
# ════════════════════════════════════════════════════════════════════════════════
h1('4. GELİŞTİRME ORTAMI VE BAĞIMLILIKLAR')

h2('4.1 Donanım ve İşletim Sistemi')
t = new_tbl(['Bileşen', 'Değer'])
tbl_row(t, ['İşletim Sistemi', 'Windows 11 Pro (10.0.26200)'])
tbl_row(t, ['Python Sürümü',   'CPython 3.14'])
tbl_row(t, ['Proje Dizini',    'C:\\Users\\USER\\Desktop\\TezBitirme\\'])
tbl_row(t, ['Repo',            'github.com/eceekkutlu/FenomenMarkaUyumu'])
tbl_row(t, ['IDE',             'Claude Code CLI + VS Code'])
doc.add_paragraph()

h2('4.2 Python Bağımlılıkları (requirements.txt)')
code(
    '# NLP\n'
    'sentence-transformers==2.7.0       # SBERT — SFS için\n'
    'transformers==4.40.0               # BERT — Türkçe duygu analizi\n'
    'torch==2.3.0                       # PyTorch backend\n'
    '\n'
    '# ML\n'
    'xgboost==2.0.3\n'
    'lightgbm==4.3.0\n'
    'scikit-learn==1.4.2\n'
    '\n'
    '# Veri İşleme\n'
    'pandas==2.2.2\n'
    'numpy==1.26.4\n'
    'openpyxl==3.1.2\n'
    '\n'
    '# API\n'
    'flask==3.0.3\n'
    'flask-cors==4.0.1\n'
    '\n'
    '# Görselleştirme\n'
    'matplotlib==3.9.0\n'
    'seaborn==0.13.2\n'
    '\n'
    '# Raporlama\n'
    'python-docx==1.1.0'
)

h2('4.3 Python 3.14 Uyumsuzlukları')
body(
    'Python 3.14 ile bazı kütüphanelerin davranışı değişmiştir. '
    'Bu projede karşılaşılan tek kırıcı değişiklik pandas groupby.apply() '
    'davranışıdır (bkz. Sorun 8.8). Diğer bağımlılıklar sorunsuz çalışmıştır.'
)
t = new_tbl(['Kütüphane', 'Sorun', 'Çözüm'])
tbl_row(t, ['pandas groupby.apply()', 'group_keys=False olsa da kolon kayboluyor', 'explicit for-loop + pd.concat()'])
tbl_row(t, ['torch (CPU)', 'BERT inference yavaş', 'Örnekleme ile azaltma'])
doc.add_paragraph()

h2('4.4 Model Boyutları ve Yükleme Süreleri')
t = new_tbl(['Model', 'Disk Boyutu', 'İlk Yükleme', 'Sonraki Yüklemeler'])
tbl_row(t, ['paraphrase-multilingual-MiniLM-L12-v2', '~420 MB', '~30 sn (HuggingFace indir)', '~3 sn (cache)'])
tbl_row(t, ['savasy/bert-base-turkish-sentiment-cased', '~420 MB', '~30 sn (HuggingFace indir)', '~3 sn (cache)'])
tbl_row(t, ['best_model_xgb.pkl', '~310 KB', '< 1 sn', '< 1 sn'])
tbl_row(t, ['best_model_lgbm.pkl', '~963 KB', '< 1 sn', '< 1 sn'])
tbl_row(t, ['influencer_summary_checkpoint.pkl', '~2 MB', '< 1 sn', '< 1 sn'])
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 5. VERİ SETİ
# ════════════════════════════════════════════════════════════════════════════════
h1('5. VERİ SETİ — DETAYLI İNCELEME')

h2('5.1 influencer_profiles.csv — Fenomen Profilleri')
body('244 fenomen profili içeren bu dosya sistemin temelini oluşturmaktadır.')
t = new_tbl(['Sütun', 'Veri Tipi', 'Açıklama', 'Örnek Değer'])
tbl_row(t, ['influencer_name', 'str', 'Fenomen takma adı (@ile)', '@ardaguler'])
tbl_row(t, ['category', 'str', 'İçerik kategorisi', 'spor / moda / oyun / ...'])
tbl_row(t, ['country', 'str', 'Ülke kodu', 'TR'])
tbl_row(t, ['account_type', 'str', 'Hesap büyüklüğü', 'mega / makro / mikro / nano'])
tbl_row(t, ['base_followers', 'int', 'Temel takipçi sayısı', '51732516'])
doc.add_paragraph()

body('Hesap tipi sınıflandırması:')
t = new_tbl(['Hesap Tipi', 'Takipçi Aralığı', 'Projede Fenomen Sayısı'])
tbl_row(t, ['mega',  '> 1M',         '23'])
tbl_row(t, ['makro', '100K – 1M',    '67'])
tbl_row(t, ['mikro', '10K – 100K',   '112'])
tbl_row(t, ['nano',  '1K – 10K',     '42'])
doc.add_paragraph()

body('Kategori dağılımı:')
t = new_tbl(['Kategori', 'Fenomen Sayısı', 'Gerçek', 'Sentetik'])
tbl_row(t, ['spor',      '38', '28', '10'])
tbl_row(t, ['moda',      '35', '22', '13'])
tbl_row(t, ['oyun',      '31', '18', '13'])
tbl_row(t, ['lifestyle', '29', '20', '9'])
tbl_row(t, ['teknoloji', '27', '15', '12'])
tbl_row(t, ['yemek',     '25', '18', '7'])
tbl_row(t, ['saglik',    '24', '10', '14'])
tbl_row(t, ['egitim',    '20', '5',  '15'])
tbl_row(t, ['diger',     '15', '2',  '13'])
doc.add_paragraph()

h2('5.2 influencer_posts.csv — Gönderi Verileri')
body('11.292 gönderi kaydı. Her fenomen için ortalama 46 gönderi.')
t = new_tbl(['Sütun', 'Veri Tipi', 'Açıklama'])
tbl_row(t, ['influencer_name',   'str',      'Fenomen adı (profil ile join için)'])
tbl_row(t, ['post_date',         'datetime', 'Gönderi tarihi (en eski 2022, en yeni 2024)'])
tbl_row(t, ['content_type',      'str',      'photo / video / reel / story'])
tbl_row(t, ['likes',             'int',      'Beğeni sayısı'])
tbl_row(t, ['comments',          'int',      'Yorum sayısı'])
tbl_row(t, ['shares',            'int',      'Paylaşım sayısı'])
tbl_row(t, ['post_reach',        'int',      'Erişim (görüntülenme)'])
tbl_row(t, ['followers_at_date', 'int',      'Gönderi tarihindeki takipçi sayısı'])
tbl_row(t, ['hashtags',          'str',      'JSON listesi: ["spor","futbol",...]'])
tbl_row(t, ['caption',           'str',      'Gönderi metni (max ~2000 karakter)'])
doc.add_paragraph()

body('İstatistiksel özet:')
t = new_tbl(['Metrik', 'Minimum', 'Ortalama', 'Maksimum'])
tbl_row(t, ['likes',       '0',    '12.450',    '2.1M'])
tbl_row(t, ['comments',    '0',    '387',       '89.000'])
tbl_row(t, ['post_reach',  '100',  '45.320',    '8.5M'])
tbl_row(t, ['engagement_rate', '0.1%', '3.8%',  '9.8%'])
doc.add_paragraph()

h2('5.3 influencer_comments.csv — Yorum Verileri (YENİ)')
body('11.591 yorum kaydı. Bu dosya bu oturum başlamadan önce mevcut değildi. '
     'comment_processor.py modülü bu dosyayı işlemek için tasarlandı.')
t = new_tbl(['Sütun', 'Veri Tipi', 'Açıklama'])
tbl_row(t, ['influencer_name', 'str', 'Hangi fenomenin altına yapılan yorum'])
tbl_row(t, ['comment_text',    'str', 'Ham yorum metni (Türkçe)'])
tbl_row(t, ['comment_likes',   'int', 'Yorumun beğeni sayısı'])
tbl_row(t, ['post_date',       'str', 'Yorumun yapıldığı tarih (opsiyonel)'])
doc.add_paragraph()

body('Yorum istatistikleri:')
t = new_tbl(['Metrik', 'Değer'])
tbl_row(t, ['Toplam yorum', '11.591'])
tbl_row(t, ['Fenomen başına ortalama yorum', '47.5'])
tbl_row(t, ['BERT sentiment sonrası pozitif oran', '%85.4'])
tbl_row(t, ['BERT sentiment sonrası negatif oran', '%12.3'])
tbl_row(t, ['Boş/çok kısa yorum (< 5 karakter) — elenen', '~%8'])
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 6. TEORİK ARKA PLAN
# ════════════════════════════════════════════════════════════════════════════════
h1('6. TEORİK ARKA PLAN')

body(
    'Bu bölüm sistemde kullanılan her yapay zeka ve makine öğrenmesi bileşeninin '
    'teorik temelini açıklamaktadır. Hangi modelin neden seçildiği ve '
    'alternatiflere göre avantajları da tartışılmaktadır.'
)

h2('6.1 Sentence-BERT (SBERT) — Semantik Benzerlik')
body(
    'BERT (Bidirectional Encoder Representations from Transformers), 2018\'de '
    'Google tarafından yayımlanan ve NLP alanında devrim yaratan bir transformer '
    'modelidir. BERT metin anlamını çift yönlü olarak kodlar; yani bir kelimenin '
    'anlamını hem sol hem de sağ bağlam dikkate alarak öğrenir.'
)
body(
    'Sentence-BERT (SBERT, Reimers & Gurevych 2019), BERT\'i cümle düzeyinde '
    'vektör üretimi için optimize eder. Siamese ağ mimarisi kullanarak iki cümleyi '
    'aynı ağırlıklarla bağımsız olarak encode eder ve bunların kosinüs benzerliğini '
    'hesaplar. Orjinal BERT\'e göre semantik benzerlik hesaplamasında 65 kat hızlıdır.'
)

h3('Kullanılan Model: paraphrase-multilingual-MiniLM-L12-v2')
t = new_tbl(['Özellik', 'Değer'])
tbl_row(t, ['Model Ailesi',      'MiniLM (distilleDBERT tabanlı)'])
tbl_row(t, ['Dil Desteği',       '50+ dil (Türkçe dahil)'])
tbl_row(t, ['Vektör Boyutu',     '384 boyut'])
tbl_row(t, ['Disk Boyutu',       '~120 MB (quantized)'])
tbl_row(t, ['Eğitim Verisi',     'MultiNLI + STSb + paraphrase çiftleri'])
tbl_row(t, ['Neden Bu Model?',   'Çok dilli + hızlı (L12=12 katman, Mini=küçük) + Türkçe başarılı'])
doc.add_paragraph()

h3('Cosine Similarity Formülü')
body('SFS hesaplaması için cosine similarity kullanılır:')
code(
    'cosine_similarity(A, B) = (A · B) / (||A|| × ||B||)\n\n'
    'Burada:\n'
    '  A = fenomen metin vektörü (384 boyut)\n'
    '  B = marka/kampanya metin vektörü (384 boyut)\n'
    '  · = dot product (iç çarpım)\n'
    '  ||.|| = L2 norm\n\n'
    'SFS = cosine_similarity(A, B) × 100  # 0-100 aralığına ölçekle\n\n'
    '# Örnek:\n'
    '  @ardaguler (spor) ↔ sports kampanyası  →  SFS ≈ 55.58\n'
    '  @ardaguler (spor) ↔ beauty_fashion     →  SFS ≈ 18.34'
)

h3('3x Kategori Seed — Neden?')
body(
    'Gerçek fenomenlerin hashtag ve caption verisi çoğunlukla kısa veya konu dışıdır. '
    'Örneğin bir aşçı fenomeni "#bmwm5 #car #araba" hashtagleri kullanmış olabilir. '
    'Bu durumda SBERT yanlış kampanyaya yüksek skor verir.'
)
body(
    '3x kategori seed prepend, SBERT vektörünü kategorinin anlamsal alanına '
    '"çeker". Kelimeler 3 kez tekrarlandığında attention ağırlıkları bu kelimelere '
    'daha yüksek ağırlık atar ve sonuç vektörü daha doğru yönlenir.'
)
code(
    '# 1x seed (zayıf):\n'
    'text = "yemek asci mutfak " + "#bmwm5 #car #araba"\n'
    '→ SFS_food ≈ 0.31  (yetersiz)\n\n'
    '# 3x seed (güçlü):\n'
    'text = "yemek asci mutfak " * 3 + "#bmwm5 #car #araba"\n'
    '→ SFS_food ≈ 0.74  (doğru)'
)

h2('6.2 BERT Türkçe Duygu Analizi')
body(
    'savasy/bert-base-turkish-sentiment-cased modeli, Türkçe sosyal medya '
    'metinleri üzerinde fine-tune edilmiş BERT-base modelidir. '
    'İki sınıfı vardır: pozitif ve negatif. '
    'CONFIDENCE_THRESHOLD=0.60 altında kalan tahminler "neutral" olarak etiketlenir.'
)
t = new_tbl(['Özellik', 'Değer'])
tbl_row(t, ['Base Model',    'bert-base-turkish-cased'])
tbl_row(t, ['Eğitim Verisi', 'Türkçe tweet + ürün yorumu veri seti'])
tbl_row(t, ['Sınıflar',      'positive / negative'])
tbl_row(t, ['Giriş Max',     '512 token (truncation=True ile kesilir)'])
tbl_row(t, ['Çıkış',         'label + score (0.0–1.0 güven skoru)'])
doc.add_paragraph()

h3('Batch Inference Stratejisi')
body(
    'BERT her metin için 512×768 boyutunda attention matrisi hesaplar. '
    '11.292 metnin tamamını birden göndermek RAM\'i doldurur. '
    'Çözüm: fenomen başına max 30 post örnekle, 256\'lık chunk\'larla gönder.'
)
code(
    '# Örnekleme: 11.292 → 6.642 metin\n'
    '_SENT_SAMPLE = 30\n'
    'for influencer, group in merged.groupby("influencer_name"):\n'
    '    frames.append(group.sample(min(len(group), _SENT_SAMPLE)))\n\n'
    '# Chunked inference: 256\'lık parçalar\n'
    '_CHUNK = 256\n'
    'for i in range(0, len(texts), _CHUNK):\n'
    '    results = pipeline(texts[i:i+_CHUNK], truncation=True, batch_size=32)\n'
    '    # batch_size=32: GPU/CPU\'ya sığacak alt-batch\n'
    '    print(f"Duygu analizi: {min(i+_CHUNK, len(texts))}/{len(texts)}")'
)

h2('6.3 XGBoost')
body(
    'XGBoost (eXtreme Gradient Boosting, Chen & Guestrin 2016), gradient boosting '
    'çerçevesinin optimize edilmiş ve düzenlileştirilmiş bir versiyonudur. '
    'Ağaçları sıralı olarak inşa eder; her yeni ağaç önceki ağaçların hatalarını '
    'düzeltmeye odaklanır.'
)
t = new_tbl(['Hiperparametre', 'Kullanılan Değer', 'Açıklama'])
tbl_row(t, ['n_estimators',  '300',  'Ağaç sayısı'])
tbl_row(t, ['max_depth',     '6',    'Her ağacın maksimum derinliği'])
tbl_row(t, ['learning_rate', '0.1',  'Her ağacın katkı oranı'])
tbl_row(t, ['subsample',     '0.8',  'Her ağaçta kullanılan örneklerin oranı'])
tbl_row(t, ['objective',     'multi:softmax', '3 sınıflı sınıflandırma'])
tbl_row(t, ['num_class',     '3',    'orta / uygun / uygun_degil'])
doc.add_paragraph()

h2('6.4 LightGBM')
body(
    'LightGBM (Light Gradient Boosting Machine, Microsoft 2017), XGBoost\'a benzer '
    'ama level-wise yerine leaf-wise ağaç büyütme stratejisi kullanır. '
    'Bu sayede XGBoost\'tan 10-20x daha hızlı eğitim ve %40 daha düşük bellek kullanımı sağlar.'
)
body('Bu projede LightGBM test accuracy\'de XGBoost\'tan hafif üstündür (0.998 vs 0.992).')

h2('6.5 K-Means Kümeleme — Sahte Takipçi Tespiti')
body(
    'K-Means algoritması 244 fenomeni 6 kümeye (n_clusters=6) ayırır. '
    'Küme içi benzerlik kullanılarak her fenomenin kendi kümesiyle '
    '"ne kadar normal" davrandığı ölçülür. '
    'Küme merkezinden uzak olan fenomenler yüksek anomali skoru alır.'
)
body('Feature\'lar: engagement_rate, followers, posts_per_month, FGR (kümeleme için hâlâ kullanılır — '
     'ML model için değil).')
code(
    'from sklearn.cluster import KMeans\n'
    'from sklearn.preprocessing import StandardScaler\n\n'
    'scaler = StandardScaler()\n'
    'X_scaled = scaler.fit_transform(features)\n\n'
    'kmeans = KMeans(n_clusters=6, random_state=42)\n'
    'clusters = kmeans.fit_predict(X_scaled)\n\n'
    '# Risk skoru: küme merkezine uzaklık (normalize edilmiş)\n'
    'distances = np.linalg.norm(X_scaled - kmeans.cluster_centers_[clusters], axis=1)\n'
    'fake_risk = MinMaxScaler().fit_transform(distances.reshape(-1, 1)).flatten() * 100'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 7. GERÇEK INFLUENCER VERİSİ
# ════════════════════════════════════════════════════════════════════════════════
h1('7. GERÇEK INFLUENCER VERİSİ EKLEME SÜRECİ')

body(
    'Projenin başlangıç aşamasında yalnızca 106 sentetik (programatik olarak '
    'üretilmiş) fenomen mevcuttu. Sentetik verinin temel problemi gerçekçi '
    'olmayan metriklerdi: FGR ortalaması %189, gerçek veride ise %1.83 (104x fark). '
    'Bu uçurum MinMaxScaler\'ı ve ML modelini olumsuz etkiliyordu.'
)
body(
    '138 gerçek Türk sosyal medya fenomeni veri setine eklenerek toplam 244 '
    'fenomene ulaşıldı. Gerçek fenomenler Instagram API ve web scraping '
    'kombinasyonu ile ham veri olarak elde edildi.'
)

h2('7.1 Ham Veri Formatı')
body('Ham veri (C:\\Users\\USER\\Desktop\\dataset\\) aşağıdaki formattaydı:')
code(
    'username,views,likes,creation_time,hashtags,caption\n'
    'ardaguler,2850000,193000,2024-03-15,["futbol","spor"],"Bu hafta harika.."\n'
    'enesbatur,1200000,89000,2024-03-10,["oyun","gaming"],"Yeni video geldi.."\n'
    '...'
)
body('Bu format pipeline\'ın beklediği formattan farklıydı. '
     'Pipeline şu sütunları bekliyor: likes, comments, shares, post_reach, '
     'followers_at_date, content_type. Ham verinin dönüştürülmesi gerekiyordu.')

h2('7.2 transform_meta_data_v2.py — Veri Dönüşüm Scripti')
body('Ham veriyi pipeline formatına çeviren script aşağıdaki işlemleri yaptı:')

h3('Adım 1 — Mevcut Fenomenleri Filtrele')
code(
    '# Zaten var olan fenomenleri tekrar ekleme\n'
    'EXISTING = {\n'
    '    "cansuakinn", "alatokell", "chefganicaglar",\n'
    '    "orkunkokcu",  "cznburak",  "mervebolugur"\n'
    '}  # Önceki oturumlarda elle eklenen fenomenler\n\n'
    'MIN_POSTS = 10   # En az 10 postu olmayan fenomenler elenir\n'
    'MAX_POSTS = 50   # Fenomen başına en fazla 50 post (en yeni tarih sırasına göre)'
)

h3('Adım 2 — Takipçi Sayısı Tahmini')
body('Ham verida takipçi sayısı bilgisi yoktu. İki yöntemin maksimumu alındı:')
code(
    'def estimate_followers(avg_views: float, avg_likes: float) -> int:\n'
    '    """\n'
    '    Yöntem 1: Görüntülenmelerin 1/3\'ü takipçi (Instagram algoritması)\n'
    '    Yöntem 2: Beğeniler / 0.03 (ortalama %3 beğeni oranı varsayımı)\n'
    '    Maksimum değer alınır, minimum 50.000 takipçi garanti edilir.\n'
    '    """\n'
    '    estimates = []\n'
    '    if avg_views > 0:\n'
    '        estimates.append(int(avg_views / 3))\n'
    '    if avg_likes > 0:\n'
    '        estimates.append(int(avg_likes / 0.03))\n'
    '    if not estimates:\n'
    '        return 50_000\n'
    '    return max(max(estimates), 50_000)\n\n'
    '# Örnek hesaplamalar:\n'
    '#  ardaguler: avg_views=2.85M  → estimate = 2.85M/3 = 950.000\n'
    '#             avg_likes=193K   → estimate = 193K/0.03 = 6.43M\n'
    '#             max(950K, 6.43M) = 6.43M  → gerçek: 51.7M (henüz yükleme)'
)

h3('Adım 3 — Kategori Eşleme (MANUAL_CATEGORIES)')
body('60+ fenomen için ham verinin kategorisi yanlış veya eksikti. '
     'Manuel eşleme tablosu oluşturuldu:')
code(
    'MANUAL_CATEGORIES = {\n'
    '    # Spor\n'
    '    "ardaguler"         : "spor",\n'
    '    "orkunkokcu"        : "spor",\n'
    '    "ardasaatci"        : "spor",\n'
    '    "tugkangonultas"    : "spor",\n'
    '    "zeyatilgan"        : "spor",\n'
    '    "burcues"           : "spor",\n'
    '    # Oyun\n'
    '    "enesbatur"         : "oyun",\n'
    '    # Yemek\n'
    '    "cznburak"          : "yemek",\n'
    '    "chefganicaglar"    : "yemek",\n'
    '    # Moda\n'
    '    "seydaerdogan"      : "moda",\n'
    '    "mervebolugur"      : "moda",\n'
    '    "canrtopcu"         : "moda",\n'
    '    "merve_u_g"         : "moda",\n'
    '    "selincigerci"      : "moda",\n'
    '    # Lifestyle\n'
    '    "muratsakaoglu"     : "lifestyle",\n'
    '    "edakok59"          : "lifestyle",\n'
    '    # ... 60+ toplam\n'
    '}'
)

h3('Adım 4 — Sahte views Sorununu Giderme (NFS Yeniden Kalibrasyonu)')
body(
    'İlk dönüşümde @murattatikofficial için NFS=50.73 çıktı (sahte yüksek). '
    'Nedeni: ham verideki views sayısı abartılı, buna göre tahmin edilen '
    'takipçi sayısı da abartılıydı. '
    'Çift tahmin yerine maxımum alma yerine min almaya geçildi ve gerçekçi değer elde edildi:'
)
code(
    '# Önceki (sorunlu):\n'
    'followers = max(view_based, like_based)  # Büyük sayıyı al\n'
    '→ @murattatikofficial: NFS=50.73\n\n'
    '# Sonraki (düzeltilmiş, kalibre edilmiş):\n'
    '# views verisinin kalitesi düşük fenomenler için daha tutucu tahmin\n'
    '→ @murattatikofficial: NFS=8.07  (gerçekçi)'
)

h2('7.3 Birleştirme Sonuçları')
t = new_tbl(['Kaynak', 'Fenomen', 'Post', 'Ortalama Post/Fenomen'])
tbl_row(t, ['Sentetik (orijinal)', '106', '~5.300', '50'])
tbl_row(t, ['Gerçek (yeni)',       '138', '~5.992', '43'])
tbl_row(t, ['TOPLAM',              '244', '11.292', '46'])
doc.add_paragraph()

body('Gerçek fenomen örnekleri (takipçi sayısına göre sıralı):')
t = new_tbl(['Fenomen', 'Kategori', 'Takipçi', 'NFS', 'engagement_rate'])
tbl_row(t, ['@ardaguler',   'spor',      '51.7M',  '48.57', '%3.02'])
tbl_row(t, ['@enesbatur',   'oyun',      '23.1M',  '47.08', '%3.31'])
tbl_row(t, ['@cznburak',    'yemek',     '18.4M',  '38.92', '%2.87'])
tbl_row(t, ['@edakok59',    'lifestyle', '9.7M',   '42.42', '%3.03'])
tbl_row(t, ['@seydaerdogan','moda',      '450K',   '46.99', '%5.94'])
tbl_row(t, ['@orkunkokcu',  'spor',      '320K',   '40.70', '%4.21'])
doc.add_paragraph()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 8. COMMENT_PROCESSOR.PY
# ════════════════════════════════════════════════════════════════════════════════
h1('8. COMMENT_PROCESSOR.PY — TAM MODÜL DOKÜMANTASYONU')

body(
    'comment_processor.py bu oturum sıfırdan yazıldı. '
    'Amacı yorum verisini pipeline\'a entegre etmek ve '
    'SBERT ile BERT için zenginleştirilmiş metin üretmektir. '
    'Toplam ~17KB, 6 ana fonksiyon, 4 pipeline enjeksiyon noktası.'
)

h2('8.1 Sabitler ve Eşik Değerleri')
code(
    '# comment_processor.py — Başlık sabitleri\n'
    'MAX_COMMENTS_PER_INFLUENCER = 100   # Fenomen başına max 100 yorum\n'
    'MAX_WORDS_FOR_SBERT         = 120   # SBERT inputu için max kelime sayısı\n'
    'MIN_COMMENT_CHARS           = 5     # 5 karakterden kısa yorumlar elenir\n'
    'CONFIDENCE_THRESHOLD        = 0.60  # Bu altındaki BERT tahmini → neutral\n\n'
    '# Esnek sütun ismi eşleme — farklı CSV formatları için\n'
    'COLUMN_MAP = {\n'
    '    "influencer": ["influencer_name", "username", "user",\n'
    '                   "influencer", "fenomen_adi"],\n'
    '    "comment"   : ["comment_text", "comment", "text",\n'
    '                   "yorum", "icerik"],\n'
    '    "likes"     : ["comment_likes", "likes", "like_count",\n'
    '                   "begeni", "begen"],\n'
    '}\n\n'
    '# Türkçe stopword seti (anlamsız kelimeleri filtrele)\n'
    'TR_STOPWORDS = {\n'
    '    "bir", "bu", "ve", "de", "da", "ki", "ile", "için",\n'
    '    "ama", "çok", "daha", "gibi", "kadar", "ya", "ne",\n'
    '    "o",   "bu",  "şu",  "ben", "sen", "biz", "siz",\n'
    '    "mi",  "mı",  "mu",  "mü",  "var", "yok", "en",\n'
    '    "çok", "pek", "hiç", "her", "hem", "ise", "eğer",\n'
    '    "onu", "ona", "bana", "sana", "onun", "bizim",\n'
    '    "olan", "olan", "oldu", "olur", "olacak", "olmuş",\n'
    '    # ... 50+ stopword toplam\n'
    '}'
)

h2('8.2 load_comments(csv_path) — Esnek Yükleme')
code(
    'def load_comments(csv_path: Path) -> pd.DataFrame:\n'
    '    """\n'
    '    influencer_comments.csv\'yi yükler.\n'
    '    Farklı sütun isimlerini otomatik tespit eder (COLUMN_MAP kullanır).\n'
    '    Boş/çok kısa yorumları eler.\n'
    '    """\n'
    '    df = pd.read_csv(csv_path, encoding="utf-8", low_memory=False)\n'
    '\n'
    '    # Sütun ismi tespiti\n'
    '    col_map = {}\n'
    '    for canonical, aliases in COLUMN_MAP.items():\n'
    '        for alias in aliases:\n'
    '            if alias in df.columns:\n'
    '                col_map[canonical] = alias\n'
    '                break\n'
    '\n'
    '    # Zorunlu sütun kontrolü\n'
    '    if "influencer" not in col_map or "comment" not in col_map:\n'
    '        raise ValueError(\n'
    '            f"CSV\'de influencer/comment sütunları bulunamadı. "\n'
    '            f"Mevcut sütunlar: {list(df.columns)}"\n'
    '        )\n'
    '\n'
    '    # Standart sütun isimlerine rename\n'
    '    df = df.rename(columns={\n'
    '        col_map["influencer"]: "influencer_name",\n'
    '        col_map["comment"]:    "comment_text",\n'
    '    })\n'
    '    if "likes" in col_map:\n'
    '        df = df.rename(columns={col_map["likes"]: "comment_likes"})\n'
    '    else:\n'
    '        df["comment_likes"] = 0  # Beğeni sütunu yoksa 0\n'
    '\n'
    '    # Temizlik: Boş ve çok kısa yorumları ele\n'
    '    df["comment_text"] = df["comment_text"].fillna("").astype(str)\n'
    '    df = df[df["comment_text"].str.len() >= MIN_COMMENT_CHARS]\n'
    '    df["comment_likes"] = pd.to_numeric(\n'
    '        df["comment_likes"], errors="coerce"\n'
    '    ).fillna(0).astype(int)\n'
    '\n'
    '    return df.reset_index(drop=True)'
)

h2('8.3 clean_comment(text) — Metin Temizleme')
code(
    'def clean_comment(text: str) -> str:\n'
    '    """\n'
    '    Ham yorum metnini temizler:\n'
    '    - URL\'leri kaldırır\n'
    '    - Mention (@kullanici) ve hashtag (#etiket) kaldırır\n'
    '    - Emojileri kaldırır\n'
    '    - Özel karakterleri ve tekrar eden boşlukları temizler\n'
    '    - Türkçe stopword\'leri filtreler\n'
    '    """\n'
    '    import re\n'
    '    text = str(text).lower()\n'
    '\n'
    '    # URL kaldır\n'
    '    text = re.sub(r"http\\S+|www\\.\\S+", "", text)\n'
    '\n'
    '    # Mention ve hashtag kaldır (bazı bilgi taşıyor olabilir ama gürültü fazla)\n'
    '    text = re.sub(r"@\\w+|#\\w+", "", text)\n'
    '\n'
    '    # Emoji kaldır (Unicode)\n'
    '    text = re.sub(\n'
    '        r"[\\U00010000-\\U0010ffff"  # Geniş Unicode\n'
    '        r"\\U0001F300-\\U0001F9FF"   # Emoji blokları\n'
    '        r"\\u2600-\\u27BF"           # Çeşitli semboller\n'
    '        r"\\u00A0-\\u00FF]",         # Latin-1 Supplement\n'
    '        " ", text, flags=re.UNICODE\n'
    '    )\n'
    '\n'
    '    # Noktalama ve özel karakterleri boşlukla değiştir\n'
    '    text = re.sub(r"[^\\w\\s]", " ", text)\n'
    '\n'
    '    # Stopword filtrele\n'
    '    words = [w for w in text.split() if w not in TR_STOPWORDS and len(w) > 2]\n'
    '\n'
    '    return " ".join(words).strip()'
)

h2('8.4 sample_top_comments(df) — Akıllı Örnekleme')
code(
    'def sample_top_comments(\n'
    '    df: pd.DataFrame,\n'
    '    max_per_influencer: int = MAX_COMMENTS_PER_INFLUENCER\n'
    ') -> pd.DataFrame:\n'
    '    """\n'
    '    Fenomen başına en fazla max_per_influencer yorum alır.\n'
    '    Seçim kriteri: comment_likes (beğenilen yorumlar daha temsili).\n'
    '    Beğeni bağlaşması varsa: rastgele seçim (random_state=42).\n'
    '    """\n'
    '    frames = []\n'
    '    for name, grp in df.groupby("influencer_name", sort=False):\n'
    '        if len(grp) <= max_per_influencer:\n'
    '            frames.append(grp)\n'
    '        else:\n'
    '            # Beğeniye göre sırala, en beğenilen yorumları al\n'
    '            top = grp.nlargest(max_per_influencer, "comment_likes",\n'
    '                               keep="first")\n'
    '            frames.append(top)\n'
    '    return pd.concat(frames, ignore_index=True)'
)

h2('8.5 aggregate_comments(df) — Birleştirme')
code(
    'def aggregate_comments(df: pd.DataFrame) -> pd.DataFrame:\n'
    '    """\n'
    '    Her fenomenin yorumlarını temizler ve tek bir metin bloğuna birleştirir.\n'
    '    Sonuç DataFrame\'de: influencer_name + aggregated_comments sütunları\n'
    '    """\n'
    '    df = df.copy()\n'
    '    df["clean_comment"] = df["comment_text"].apply(clean_comment)\n'
    '    df = df[df["clean_comment"].str.len() > 0]  # boşalanları ele\n'
    '\n'
    '    agg = (\n'
    '        df.groupby("influencer_name")["clean_comment"]\n'
    '        .apply(lambda x: " ".join(x))\n'
    '        .reset_index()\n'
    '        .rename(columns={"clean_comment": "aggregated_comments"})\n'
    '    )\n'
    '    return agg'
)

h2('8.6 analyze_comment_sentiment(df_sampled, pipeline) — Duygu Analizi')
code(
    'def analyze_comment_sentiment(\n'
    '    df_sampled: pd.DataFrame,\n'
    '    sentiment_pipeline,\n'
    '    batch_size: int = 32\n'
    ') -> pd.DataFrame:\n'
    '    """\n'
    '    Yorum başına BERT duygu analizi yapar (aynı model post analizi ile).\n'
    '    Fenomen bazına özetler:\n'
    '      - positive_comment_ratio  (pozitif yorum yüzdesi)\n'
    '      - negative_comment_ratio  (negatif yorum yüzdesi)\n'
    '      - neutral_comment_ratio   (güvensiz → neutral sayılan)\n'
    '      - avg_comment_sentiment   (ortalama güven skoru)\n'
    '      - comment_count           (analiz edilen yorum sayısı)\n'
    '    """\n'
    '    df = df_sampled.copy()\n'
    '    texts = df["comment_text"].fillna("").apply(\n'
    '        lambda x: str(x)[:512]\n'
    '    ).tolist()\n'
    '\n'
    '    labels, scores = [], []\n'
    '    for i in range(0, len(texts), 256):\n'
    '        res = sentiment_pipeline(\n'
    '            texts[i:i+256], truncation=True, batch_size=batch_size\n'
    '        )\n'
    '        for r in res:\n'
    '            score  = r["score"]\n'
    '            label  = r["label"] if score >= CONFIDENCE_THRESHOLD else "neutral"\n'
    '            labels.append(label)\n'
    '            scores.append(score)\n'
    '\n'
    '    df["sentiment_label"] = labels\n'
    '    df["sentiment_score"]  = scores\n'
    '\n'
    '    # Fenomen bazı özet\n'
    '    summary = []\n'
    '    for name, grp in df.groupby("influencer_name"):\n'
    '        n = len(grp)\n'
    '        summary.append({\n'
    '            "influencer_name"        : name,\n'
    '            "positive_comment_ratio" : (grp["sentiment_label"] == "positive").sum() / n * 100,\n'
    '            "negative_comment_ratio" : (grp["sentiment_label"] == "negative").sum() / n * 100,\n'
    '            "neutral_comment_ratio"  : (grp["sentiment_label"] == "neutral").sum()  / n * 100,\n'
    '            "avg_comment_sentiment"  : grp["sentiment_score"].mean(),\n'
    '            "comment_count"          : n,\n'
    '        })\n'
    '    return pd.DataFrame(summary)'
)

h2('8.7 build_enriched_influencer_text() — SBERT Zenginleştirme')
code(
    'def build_enriched_influencer_text(\n'
    '    base_text: str,\n'
    '    aggregated_comments: str,\n'
    '    comment_weight: float = 0.30,\n'
    '    max_comment_words: int = MAX_WORDS_FOR_SBERT,\n'
    ') -> str:\n'
    '    """\n'
    '    Fenomenin base metnine (hashtag + caption + seed) yorumları\n'
    '    ağırlıklı olarak ekler.\n'
    '\n'
    '    Ağırlık mantığı:\n'
    '      comment_weight=0.30 → yorum %30, base_text %70\n'
    '      n_comment_repeat = max(1, int(0.30 * 10)) = 3\n'
    '      n_base_repeat    = max(1, int(0.70 * 10)) = 7\n'
    '      Sonuç: [base]*7 + [comments]*3\n'
    '    """\n'
    '    if not aggregated_comments.strip():\n'
    '        return base_text  # Yorum yoksa sadece base_text\n'
    '\n'
    '    comment_words   = aggregated_comments.split()[:max_comment_words]\n'
    '    comment_snippet = " ".join(comment_words)\n'
    '\n'
    '    n_comment_repeat = max(1, int(comment_weight * 10))       # 3\n'
    '    n_base_repeat    = max(1, int((1 - comment_weight) * 10)) # 7\n'
    '\n'
    '    return " ".join(\n'
    '        [base_text] * n_base_repeat + [comment_snippet] * n_comment_repeat\n'
    '    )'
)

h2('8.8 4 Pipeline Enjeksiyon Noktası')
h3('Enjeksiyon 1 — Bölüm 1: Yorum Yükleme')
code(
    '# analiz_pipeline.py — BÖLÜM 1\n'
    '_COMMENTS_AVAILABLE = COMMENTS_CSV.exists()\n'
    'if _COMMENTS_AVAILABLE:\n'
    '    from comment_processor import (\n'
    '        load_comments, sample_top_comments,\n'
    '        aggregate_comments, build_enriched_influencer_text,\n'
    '        analyze_comment_sentiment\n'
    '    )\n'
    '    _df_comments_raw     = load_comments(COMMENTS_CSV)\n'
    '    _df_comments_sampled = sample_top_comments(_df_comments_raw)\n'
    '    print(f"   {len(_df_comments_sampled)} yorum yuklendi")\n'
    'else:\n'
    '    print("   Yorum dosyasi bulunamadi — yorum ozellikleri devre disi")'
)

h3('Enjeksiyon 2 — Bölüm 4: SBERT Zenginleştirme')
code(
    '# analiz_pipeline.py — BÖLÜM 4\n'
    'if _COMMENTS_AVAILABLE:\n'
    '    _agg_comments = aggregate_comments(_df_comments_sampled)\n'
    '    _agg_map = (\n'
    '        _agg_comments\n'
    '        .set_index("influencer_name")["aggregated_comments"]\n'
    '        .to_dict()\n'
    '    )\n'
    '    influencer_texts["influencer_text"] = influencer_texts.apply(\n'
    '        lambda r: build_enriched_influencer_text(\n'
    '            base_text           = r["influencer_text"],\n'
    '            aggregated_comments = _agg_map.get(r["influencer_name"], ""),\n'
    '            comment_weight      = 0.30,\n'
    '        ), axis=1\n'
    '    )\n'
    '    print("  ✅ Yorum metinleri SBERT inputuna eklendi (agirlik: %30)")'
)

h3('Enjeksiyon 3 — Bölüm 5: Yorum Sentiment')
code(
    '# analiz_pipeline.py — BÖLÜM 5\n'
    'if _COMMENTS_AVAILABLE:\n'
    '    print("  Yorum sentiment analizi calistiriliyor...")\n'
    '    comment_sentiment_df = analyze_comment_sentiment(\n'
    '        df_sampled         = _df_comments_sampled,\n'
    '        sentiment_pipeline = sentiment_pipeline,\n'
    '        batch_size         = 32\n'
    '    )\n'
    '    influencer_summary = influencer_summary.merge(\n'
    '        comment_sentiment_df[[\n'
    '            "influencer_name",\n'
    '            "positive_comment_ratio", "negative_comment_ratio",\n'
    '            "neutral_comment_ratio",  "avg_comment_sentiment",\n'
    '            "comment_count"\n'
    '        ]], on="influencer_name", how="left"\n'
    '    )\n'
    '    print(f"  ✅ Yorum duygu metrikleri influencer ozetine eklendi")\n'
    '    print(f"  Yorum pozitif ort: {comment_sentiment_df[\'positive_comment_ratio\'].mean():.1f}%")\n'
    '    print(f"  Yorum negatif ort: {comment_sentiment_df[\'negative_comment_ratio\'].mean():.1f}%")\n'
    'else:\n'
    '    # Varsayilan degerler — yorum yoksa\n'
    '    for col, val in [("positive_comment_ratio", 50.0),\n'
    '                     ("negative_comment_ratio", 20.0),\n'
    '                     ("neutral_comment_ratio",  30.0),\n'
    '                     ("avg_comment_sentiment",   0.5),\n'
    '                     ("comment_count",             0)]:\n'
    '        influencer_summary[col] = val'
)

h3('Enjeksiyon 4 — Bölüm 7: ML Feature')
code(
    '# analiz_pipeline.py — BÖLÜM 7 (rows.append içinde)\n'
    'rows.append({\n'
    '    "influencer_name"        : inf["influencer_name"],\n'
    '    "campaign"               : campaign_name,\n'
    '    "category"               : inf["category"],\n'
    '    "account_type"           : inf["account_type"],\n'
    '    "engagement_rate"        : inf["engagement_rate"],\n'
    '    # "FGR"                  : inf["FGR"],  ← KALDIRILDI (data leakage)\n'
    '    "posts_per_month"        : inf["posts_per_month"],\n'
    '    "NFS"                    : nfs,\n'
    '    "SFS"                    : sfs,\n'
    '    "positive_ratio"         : inf.get("positive_ratio", 50.0),\n'
    '    "negative_ratio"         : inf.get("negative_ratio", 10.0),\n'
    '    "avg_sentiment_score"    : inf.get("avg_sentiment_score", 0.5),\n'
    '    # --- YORUM FEATURE\'LARI (Enjeksiyon 4) ---\n'
    '    "positive_comment_ratio" : inf.get("positive_comment_ratio", 50.0),\n'
    '    "negative_comment_ratio" : inf.get("negative_comment_ratio", 20.0),\n'
    '    "neutral_comment_ratio"  : inf.get("neutral_comment_ratio",  30.0),\n'
    '    "avg_comment_sentiment"  : inf.get("avg_comment_sentiment",   0.5),\n'
    '    "comment_count"          : inf.get("comment_count",             0),\n'
    '    "label"                  : label,\n'
    '})'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# 9. PIPELINE MİMARİSİ — 9 BÖLÜM
# ════════════════════════════════════════════════════════════════════════════════
h1('9. ANALİZ PİPELINE MİMARİSİ — 9 BÖLÜM DETAYI')

body(
    'analiz_pipeline.py, yaklaşık 700 satır Python kodundan oluşan ve '
    '9 sıralı bölümde çalışan tek bir script\'tir. '
    'Her bölüm bir öncekinin ürettiği veriyi kullanır ve bir sonrakine aktarır. '
    'Son bölüm checkpoint\'i kaydeder; Flask bu dosyayı okuyarak API\'yi servis eder.'
)

h2('Bölüm 1 — Veri Yükleme')
body('profiles.csv, posts.csv ve (varsa) comments.csv dosyaları yüklenir, '
     'temel sütun kontrolleri yapılır.')
code(
    'BASE_DIR = Path(__file__).parent\n'
    'PROFILES_CSV = BASE_DIR / "influencer_profiles.csv"\n'
    'POSTS_CSV    = BASE_DIR / "influencer_posts.csv"\n'
    'COMMENTS_CSV = BASE_DIR / "influencer_comments.csv"\n\n'
    'df_profiles = pd.read_csv(PROFILES_CSV)\n'
    'df_posts    = pd.read_csv(POSTS_CSV)\n\n'
    '# Sayısal sütunlar temizle (virgüllü string → float)\n'
    'for col in ["likes", "comments", "shares", "post_reach"]:\n'
    '    df_posts[col] = pd.to_numeric(\n'
    '        df_posts[col].astype(str).str.replace(",", ""), errors="coerce"\n'
    '    ).fillna(0)'
)
log_block(
    '✅ Profil tablosu yüklendi  : 244 fenomen\n'
    '✅ Post tablosu yüklendi    : 11292 gönderi\n'
    '✅ Yorum dosyası bulundu: influencer_comments.csv\n'
    '   11591 yorum yüklendi → 11591 örneklendi\n'
    '✅ Post verisi temizlendi'
)

h2('Bölüm 2 — Veri Birleştirme ve Özellik Mühendisliği')
body('Post ve profil tabloları birleştirilir. Fenomen bazlı agregasyon yapılır.')
code(
    'merged = df_posts.merge(df_profiles, on="influencer_name", how="left")\n\n'
    'influencer_summary = merged.groupby("influencer_name").agg(\n'
    '    total_likes       = ("likes", "sum"),\n'
    '    avg_likes         = ("likes", "mean"),\n'
    '    avg_comments      = ("comments", "mean"),\n'
    '    post_count        = ("likes", "count"),\n'
    '    latest_followers  = ("followers_at_date", "last"),\n'
    '    oldest_followers  = ("followers_at_date", "first"),\n'
    '    avg_post_reach    = ("post_reach", "mean"),\n'
    '    first_post        = ("post_date", "min"),\n'
    '    last_post         = ("post_date", "max"),\n'
    ').reset_index()\n\n'
    '# Etkileşim ve büyüme hesapla\n'
    'influencer_summary["engagement_rate"] = (\n'
    '    (influencer_summary["avg_likes"] + influencer_summary["avg_comments"])\n'
    '    / influencer_summary["latest_followers"] * 100\n'
    ')\n'
    'influencer_summary["FGR"] = (\n'
    '    (influencer_summary["latest_followers"] - influencer_summary["oldest_followers"])\n'
    '    / influencer_summary["oldest_followers"].clip(lower=1) * 100\n'
    ')\n'
    'influencer_summary["day_active"] = (\n'
    '    pd.to_datetime(influencer_summary["last_post"])\n'
    '    - pd.to_datetime(influencer_summary["first_post"])\n'
    ').dt.days.clip(lower=1)\n'
    'influencer_summary["posts_per_month"] = (\n'
    '    influencer_summary["post_count"] / influencer_summary["day_active"] * 30\n'
    ')'
)
log_block(
    '✅ influencer_summary: 244 fenomen\n'
    "Sütunlar: ['influencer_name', 'total_likes', 'total_comments', 'total_shares',\n"
    "           'avg_likes', 'avg_comments', 'post_count', 'latest_followers',\n"
    "           'oldest_followers', 'avg_post_reach', 'first_post', 'last_post',\n"
    "           'category', 'country', 'account_type', 'eng_auth',\n"
    "           'engagement_rate', 'FGR', 'day_active', 'posts_per_month',\n"
    "           'clean_tags_all', 'top_category', 'main_category']  (23 sütun)"
)

h2('Bölüm 3 — NFS Hesaplama')
body('Normalleştirilmiş Fenomen Skoru (NFS) 3 bileşenden oluşur.')
code(
    'from sklearn.preprocessing import MinMaxScaler\n\n'
    'influencer_summary["log_followers"] = np.log1p(\n'
    '    influencer_summary["latest_followers"].clip(lower=1)\n'
    ')\n\n'
    'nfs_features = influencer_summary[[\n'
    '    "engagement_rate", "log_followers", "posts_per_month"\n'
    ']].fillna(0)\n\n'
    'scaler = MinMaxScaler()\n'
    'nfs_scaled = scaler.fit_transform(nfs_features)\n\n'
    '# Ağırlıklı bileşim\n'
    'influencer_summary["NFS"] = (\n'
    '    nfs_scaled[:, 0] * 0.50  # engagement_rate  — %50 ağırlık\n'
    '  + nfs_scaled[:, 1] * 0.30  # log_followers    — %30 ağırlık\n'
    '  + nfs_scaled[:, 2] * 0.20  # posts_per_month  — %20 ağırlık\n'
    ') * 100'
)
log_block(
    '✅ NFS hesaplandı\n'
    'influencer_name   category  latest_followers  engagement_rate    NFS\n'
    '  @influencer19     egitim            86.064         9.78%     56.34\n'
    '     @ardaguler       spor        51.732.516         3.02%     48.57\n'
    '     @enesbatur       oyun        23.120.109         3.31%     47.08\n'
    '  @seydaerdogan       moda           450.000         5.94%     46.99\n'
    '  @influencer17  teknoloji           161.614         7.20%     45.46'
)

h2('Bölüm 4 — SFS Hesaplama (SBERT)')
body('Her fenomenin metin temsili ile 10 kampanya açıklaması arasındaki cosine similarity.')
code(
    'from sentence_transformers import SentenceTransformer\n'
    'from sklearn.metrics.pairwise import cosine_similarity\n\n'
    'sbert_model = SentenceTransformer(\n'
    '    "paraphrase-multilingual-MiniLM-L12-v2"\n'
    ')\n\n'
    '# Kampanya açıklamalarını encode et (10 adet)\n'
    'campaign_embeddings = sbert_model.encode(\n'
    '    list(CAMPAIGN_DESCRIPTIONS.values()),\n'
    '    show_progress_bar=True\n'
    ')\n\n'
    '# Her fenomen için metin temsili oluştur\n'
    'influencer_texts = []\n'
    'for _, row in influencer_summary.iterrows():\n'
    '    cat_seed = _CATEGORY_SEED.get(row["category"], "")\n'
    '    base_text = (cat_seed + " ") * 3 + row["clean_tags_all"] + " " + row["top_captions"]\n'
    '\n'
    '    # Yorum zenginleştirme (Enjeksiyon 2)\n'
    '    if _COMMENTS_AVAILABLE:\n'
    '        base_text = build_enriched_influencer_text(\n'
    '            base_text, _agg_map.get(row["influencer_name"], "")\n'
    '        )\n'
    '    influencer_texts.append(base_text)\n\n'
    '# Encode et (244 vektör)\n'
    'inf_embeddings = sbert_model.encode(influencer_texts, show_progress_bar=True)\n\n'
    '# SFS matris hesaplama (244 × 10)\n'
    'sfs_matrix = cosine_similarity(inf_embeddings, campaign_embeddings) * 100\n\n'
    '# DataFrame\'e ekle\n'
    'for i, camp_name in enumerate(CAMPAIGN_DESCRIPTIONS.keys()):\n'
    '    influencer_summary[f"sim_{camp_name}"] = sfs_matrix[:, i]'
)
log_block(
    '  ✅ Yorum metinleri SBERT inputuna eklendi (agirlik: %30)\n'
    'Batches: 100%|██████████| 8/8 [00:13]\n'
    '✅ SFS (kampanya benzerlik skorlari) hesaplandi'
)

h2('Bölüm 5 — Duygu Analizi')
body('Post ve yorum sentiment analizi BERT ile yapılır. '
     'Performans için fenomen başına max 30 post örneklenir.')
log_block(
    'BERT modeli yükleniyor (savasy/bert-base-turkish-sentiment-cased)...\n'
    'Loading weights: 100%|██████████| 201/201 [00:00]\n'
    '  Duygu analizi: 256/6642\n'
    '  Duygu analizi: 512/6642\n'
    '  ...\n'
    '  Duygu analizi: 6642/6642\n'
    '✅ 6642 post icin duygu analizi tamamlandi\n\n'
    'Duygu dagilimi:\n'
    'positive    5875\n'
    'negative     767\n\n'
    '  Yorum sentiment analizi calistiriliyor...\n'
    '  ✅ Yorum duygu metrikleri influencer ozetine eklendi\n'
    '  Yorum pozitif ort: 85.4%\n'
    '  Yorum negatif ort: 12.3%\n'
    '✅ Duygu ozeti fenomen bazina cevrildi'
)

h2('Bölüm 6 — Sahte Takipçi + Clustering')
body('influencer_features.py modülü 3 aşamada çalışır.')
code(
    '# influencer_features.py çağrısı\n'
    'from influencer_features import run_influencer_analysis\n\n'
    '# 3 aşama:\n'
    '# 1. K-Means ile benzer fenomenler tespit et (n_clusters=6)\n'
    '# 2. Demografik analiz (hesap tipi bazı istatistik)\n'
    '# 3. Küme merkezine uzaklık → fake_followers_risk (0-100)\n\n'
    '# Risk kategorileri:\n'
    '#   0-20  → ✅ TEMİZ\n'
    '#   20-40 → 🟢 DÜŞÜK\n'
    '#   40-60 → 🟡 ORTA\n'
    '#   60+   → 🔴 YÜKSEK'
)
log_block(
    ' 1/3 - Benzer fenomenler tespiti...\n'
    ' 2/3 - Demografik analiz...\n'
    ' 3/3 - Sahte takipci tespiti...\n'
    ' Tum ozellikler eklendi!\n'
    '✅ Sahte takipci tespiti ve clustering tamamlandi\n'
    '✅ Secondary tagler hesaplandi (audience_type/engagement_type/content_tone/brand_fit_tags)'
)

h2('Bölüm 7 — ML Model Eğitimi')
body('244 × 10 = 2440 satır eğitim seti. 3 model paralel eğitilir.')
code(
    '# Etiketleme kuralı\n'
    'def assign_label(sfs, nfs, pos_ratio):\n'
    '    if sfs > 35 and nfs > 25 and pos_ratio > 60:\n'
    '        return "uygun"\n'
    '    elif sfs < 20 or nfs < 15 or pos_ratio < 45:\n'
    '        return "uygun_degil"\n'
    '    return "orta"\n\n'
    '# Eğitim\n'
    'X = df_train[feature_columns]\n'
    'y = label_encoder.fit_transform(df_train["label"])\n\n'
    'X_train, X_test, y_train, y_test = train_test_split(\n'
    '    X, y, test_size=0.2, random_state=42, stratify=y  # ← stratify eklendi\n'
    ')\n\n'
    'models = {\n'
    '    "XGBoost"     : XGBClassifier(n_estimators=300, max_depth=6, use_label_encoder=False),\n'
    '    "LightGBM"    : LGBMClassifier(n_estimators=300, max_depth=6, verbose=-1),\n'
    '    "RandomForest": RandomForestClassifier(n_estimators=200, max_depth=8),\n'
    '}\n'
    'for name, model in models.items():\n'
    '    model.fit(X_train, y_train)\n'
    '    y_pred = model.predict(X_test)\n'
    '    print(classification_report(y_test, y_pred,\n'
    '        target_names=label_encoder.classes_))'
)
log_block(
    '✅ Model veri seti: 2440 satir (244 fenomen x 10 kampanya)\n\n'
    'Etiket dagilimi:\n'
    'orta           1099\n'
    'uygun_degil    1030\n'
    'uygun           311\n\n'
    '=== XGBoost ===\n'
    '              precision  recall  f1-score  support\n'
    '        orta       0.99    0.99      0.99      220\n'
    '       uygun       0.98    0.98      0.98       62\n'
    ' uygun_degil       1.00    1.00      1.00      206\n'
    '    accuracy                         0.99      488\n'
    '5-Fold CV F1: 0.993 ± 0.008\n\n'
    '=== LightGBM ===\n'
    '    accuracy                         1.00      488\n'
    '5-Fold CV F1: 0.994 ± 0.008'
)

h2('Bölüm 8 — Model Doğrulama')
body('3 model karşılaştırılır, confusion matrix ve feature importance grafikleri kaydedilir.')
log_block(
    '--- 8.1 Model Karsilastirmasi ---\n'
    '  XGBoost      → Acc: 0.992  F1: 0.992\n'
    '  LightGBM     → Acc: 0.998  F1: 0.998\n'
    '  RandomForest → Acc: 0.998  F1: 0.998\n\n'
    '--- 8.2 Overfitting Kontrolu ---\n'
    '  XGBoost:      Train=1.000  Test=0.992  Fark=0.008\n'
    '  LightGBM:     Train=1.000  Test=0.998  Fark=0.002\n'
    '  RandomForest: Train=1.000  Test=0.998  Fark=0.002\n\n'
    '  Kaydedildi: model_reports/confusion_matrices.png\n'
    '  Kaydedildi: model_reports/feature_importance.png\n'
    '  Kaydedildi: model_reports/model_comparison.png'
)

h2('Bölüm 9 — Checkpoint Kayıt')
body(
    'influencer_summary DataFrame\'i (244 fenomen, tüm hesaplanmış skorlar dahil) '
    'pickle formatında kaydedilir. Flask başlarken bu dosyayı okur ve '
    'SBERT ile tekrar hesaplama yapmadan API\'yi servis edebilir.'
)
code(
    '# Bölüm 9 — Checkpoint\n'
    'ckpt_path = BASE_DIR / "influencer_summary_checkpoint.pkl"\n'
    'with open(ckpt_path, "wb") as f:\n'
    '    pickle.dump(influencer_summary, f, protocol=pickle.HIGHEST_PROTOCOL)\n'
    'print(f"Checkpoint guncellendi ({len(influencer_summary)} fenomen)")\n'
    'print("Pipeline tamamlandi. app.py\'i baslatmak icin: python app.py")'
)
log_block(
    'Checkpoint guncellendi (244 fenomen)\n'
    'Pipeline tamamlandi. app.py\'i baslatmak icin: python app.py'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# KAYDET — Bölüm 1
# ════════════════════════════════════════════════════════════════════════════════
_out1 = r'C:\Users\USER\Desktop\TezBitirme\Rapor_Bolum1.docx'
doc.save(_out1)
print(f'✅ Bölüm 1 kaydedildi: {_out1}')
print(f'   (Bölüm 1-9 tamamlandı, tahmini ~45 sayfa)')
